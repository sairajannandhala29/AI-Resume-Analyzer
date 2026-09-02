import os
import re
import inspect
import tempfile
import textwrap
from pathlib import Path

import streamlit as st
from docx import Document
from dotenv import load_dotenv

load_dotenv()

# ============================================================
# OPTIONAL / COMPATIBLE MODULE IMPORTS
# ============================================================

def _load_callable(candidates, required=True):
    """
    Load a callable from the first matching module/function pair.
    This keeps app.py compatible with the existing project modules.
    """
    import importlib

    errors = []

    for module_name, function_names in candidates:
        try:
            module = importlib.import_module(module_name)
        except Exception as exc:
            errors.append(f"{module_name}: {exc}")
            continue

        for function_name in function_names:
            fn = getattr(module, function_name, None)
            if callable(fn):
                return fn

    if required:
        attempted = ", ".join(
            f"{m}.{'/'.join(fs)}" for m, fs in candidates
        )
        raise ImportError(
            "Could not load required project function. "
            f"Attempted: {attempted}"
        )

    return None


# Core project functions. The candidate list allows the app to work
# with the module names used in the existing project.
extract_text_from_pdf = _load_callable(
    [
        (
            "modules.pdf_parser",
            ["extract_text_from_pdf"],
        ),
        (
            "modules.resume_parser",
            ["extract_text_from_pdf"],
        ),
        (
            "modules.parser",
            ["extract_text_from_pdf"],
        ),
    ],
    required=False,
)

clean_text = _load_callable(
    [
        ("modules.text_cleaner", ["clean_text"]),
        ("modules.text_processing", ["clean_text"]),
        ("modules.text_utils", ["clean_text"]),
    ],
    required=False,
)

extract_skills = _load_callable(
    [
        ("modules.skill_extractor", ["extract_skills"]),
        ("modules.skills", ["extract_skills"]),
        ("modules.jd_analyzer", ["extract_skills"]),
    ]
)

analyze_job_description = _load_callable(
    [
        ("modules.jd_analyzer", ["analyze_job_description"]),
        ("modules.job_analyzer", ["analyze_job_description"]),
        ("modules.jd_parser", ["analyze_job_description"]),
    ]
)

calculate_ats_score = _load_callable(
    [
        ("modules.ats_scorer", ["calculate_ats_score"]),
        ("modules.ats_score", ["calculate_ats_score"]),
        ("modules.scoring", ["calculate_ats_score"]),
    ]
)

generate_recommendations = _load_callable(
    [
        ("modules.recommendations", ["generate_recommendations"]),
        ("modules.recommendations", ["get_recommendations"]),
        ("modules.ats_recommendations", ["generate_recommendations"]),
        ("modules.ats_recommendations", ["get_recommendations"]),
    ],
    required=False,
)

identify_strengths = _load_callable(
    [
        ("modules.strengths", ["identify_strengths"]),
        ("modules.strengths", ["get_strengths"]),
        ("modules.ats_strengths", ["identify_strengths"]),
        ("modules.ats_strengths", ["get_strengths"]),
    ],
    required=False,
)

# Optimization functions are known to live in resume_optimizer.py.
generate_verified_optimized_resume = _load_callable(
    [
        (
            "modules.resume_optimizer",
            ["generate_verified_optimized_resume"],
        ),
    ]
)

from modules.resume_editor import update_existing_resume

get_available_providers = _load_callable(
    [
        ("modules.ai_rewriter", ["get_available_providers"]),
        ("modules.resume_optimizer", ["get_available_providers"]),
    ],
    required=False,
)


# ============================================================
# FALLBACK HELPERS
# ============================================================

def fallback_clean_text(text):
    text = text or ""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_text(text):
    if clean_text:
        try:
            result = clean_text(text)
            if result:
                return result
        except Exception:
            pass
    return fallback_clean_text(text)


def fallback_extract_pdf_text(file_bytes):
    """
    Fallback PDF extraction if the project's PDF parser cannot be loaded.
    """
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError(
            "PyMuPDF is required to read PDF resumes."
        ) from exc

    document = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []

    try:
        for page in document:
            text = page.get_text("text")
            if text:
                pages.append(text)
    finally:
        document.close()

    return "\n".join(pages).strip()


def read_uploaded_resume(uploaded_file):
    """
    Read PDF or DOCX into plain text.
    """
    suffix = Path(uploaded_file.name).suffix.lower()
    file_bytes = uploaded_file.getvalue()

    if suffix == ".pdf":
        if extract_text_from_pdf:
            # Existing parser may accept bytes or a path. Try bytes first.
            try:
                text = extract_text_from_pdf(file_bytes)
                if text:
                    return normalize_text(text)
            except Exception:
                pass

            # Fall back to a temporary file path.
            with tempfile.NamedTemporaryFile(
                suffix=".pdf",
                delete=False,
            ) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            try:
                text = extract_text_from_pdf(tmp_path)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

            return normalize_text(text)

        return normalize_text(
            fallback_extract_pdf_text(file_bytes)
        )

    if suffix == ".docx":
        document = Document(
            _write_temp_file(file_bytes, ".docx")
        )

        try:
            parts = []

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(text)

            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        parts.append(" | ".join(row_text))

            return normalize_text("\n".join(parts))
        finally:
            # Document opens from the temp path created above.
            pass

    raise ValueError(
        "Unsupported resume format. Please upload a PDF or DOCX."
    )


def _write_temp_file(file_bytes, suffix):
    tmp = tempfile.NamedTemporaryFile(
        suffix=suffix,
        delete=False,
    )
    try:
        tmp.write(file_bytes)
        tmp.flush()
        return tmp.name
    finally:
        tmp.close()


def extract_docx_text(docx_path):
    """
    Extract text from paragraphs and tables of the generated DOCX.
    This is the text that is actually re-scored.
    """
    document = Document(docx_path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    row_text.append(text)
            if row_text:
                parts.append(" | ".join(row_text))

    return normalize_text("\n".join(parts))


def _get_first(data, keys, default=None):
    """
    Read a value from a dict using several likely keys.
    """
    if not isinstance(data, dict):
        return default

    for key in keys:
        if key in data:
            return data[key]

    return default


def extract_numeric_score(value):
    """
    Convert the different possible score return shapes into one float.
    Supports:
      34.59
      0.3459
      {"ats_score": 34.59}
      {"score": 34.59}
      nested dicts
    """
    if isinstance(value, bool):
        return None

    if isinstance(value, (int, float)):
        number = float(value)
        if 0 <= number <= 1:
            return number * 100
        return number

    if isinstance(value, str):
        match = re.search(
            r"(-?\d+(?:\.\d+)?)\s*%",
            value,
        )
        if match:
            return float(match.group(1))

        try:
            number = float(value.strip())
            if 0 <= number <= 1:
                return number * 100
            return number
        except ValueError:
            return None

    if isinstance(value, dict):
        preferred = [
            "ats_score",
            "overall_score",
            "overall",
            "total_score",
            "score",
            "final_score",
            "ats",
        ]

        for key in preferred:
            if key in value:
                result = extract_numeric_score(value[key])
                if result is not None:
                    return result

        # Search one level deeper.
        for nested in value.values():
            result = extract_numeric_score(nested)
            if result is not None:
                return result

    return None


def get_score_from_analysis(analysis):
    score = extract_numeric_score(analysis)
    if score is None:
        return 0.0
    return max(0.0, min(100.0, score))


def get_list_from_dict(data, keys):
    if not isinstance(data, dict):
        return []

    for key in keys:
        value = data.get(key)

        if isinstance(value, (list, tuple, set)):
            return [
                str(item).strip()
                for item in value
                if str(item).strip()
            ]

    return []


def get_resume_skills(analysis):
    return get_list_from_dict(
        analysis,
        [
            "resume_skills",
            "skills",
            "extracted_skills",
            "found_skills",
        ],
    )


def get_jd_skills(jd_analysis):
    return get_list_from_dict(
        jd_analysis,
        [
            "jd_skills",
            "skills",
            "required_skills",
            "extracted_skills",
        ],
    )


def call_ats_scorer(
    resume_text,
    job_description,
    resume_skills,
    jd_analysis,
):
    """
    Call the existing ATS scorer while tolerating common signatures.
    """
    attempts = [
        (
            resume_text,
            job_description,
            resume_skills,
            jd_analysis,
        ),
        (
            resume_text,
            job_description,
            resume_skills,
        ),
        (
            resume_text,
            job_description,
        ),
        (
            resume_text,
            jd_analysis,
        ),
    ]

    # Prefer a signature-aware keyword call when possible.
    try:
        signature = inspect.signature(calculate_ats_score)
        params = list(signature.parameters.values())

        kwargs = {}

        for parameter in params:
            name = parameter.name.lower()

            if "resume" in name and "skill" in name:
                kwargs[parameter.name] = resume_skills
            elif (
                "resume" in name
                and ("text" in name or "content" in name)
            ):
                kwargs[parameter.name] = resume_text
            elif (
                "job" in name
                or "jd" in name
                or "description" in name
            ):
                if "analysis" in name:
                    kwargs[parameter.name] = jd_analysis
                else:
                    kwargs[parameter.name] = job_description
            elif "analysis" in name:
                kwargs[parameter.name] = jd_analysis

        required_params = [
            p
            for p in params
            if p.default is inspect.Parameter.empty
            and p.kind
            in (
                inspect.Parameter.POSITIONAL_ONLY,
                inspect.Parameter.POSITIONAL_OR_KEYWORD,
            )
        ]

        if len(kwargs) >= len(required_params):
            try:
                return calculate_ats_score(**kwargs)
            except TypeError:
                pass

    except Exception:
        pass

    last_error = None

    for args in attempts:
        try:
            return calculate_ats_score(*args)
        except TypeError as exc:
            last_error = exc
            continue

    raise TypeError(
        "Could not call calculate_ats_score with the existing "
        "project signature."
    ) from last_error


def run_ats_analysis(resume_text, job_description):
    """
    Run the same ATS pipeline for both original and optimized resumes.
    """
    resume_text = normalize_text(resume_text)
    job_description = normalize_text(job_description)

    resume_skills = extract_skills(resume_text)
    if resume_skills is None:
        resume_skills = []

    jd_analysis = analyze_job_description(job_description)
    score_result = call_ats_scorer(
        resume_text,
        job_description,
        resume_skills,
        jd_analysis,
    )

    return {
        "ats_score": get_score_from_analysis(score_result),
        "score_result": score_result,
        "resume_skills": list(resume_skills),
        "jd_analysis": jd_analysis,
        "jd_skills": get_jd_skills(jd_analysis),
    }


def safe_recommendations(ats_analysis):
    if not generate_recommendations:
        return []

    candidates = [
        (ats_analysis,),
        (
            ats_analysis.get("score_result"),
            ats_analysis.get("jd_analysis"),
        ),
        (
            ats_analysis.get("score_result"),
        ),
    ]

    for args in candidates:
        try:
            result = generate_recommendations(*args)

            if isinstance(result, str):
                return [result]

            if isinstance(result, (list, tuple)):
                return [str(x) for x in result]

            if isinstance(result, dict):
                values = (
                    result.get("recommendations")
                    or result.get("items")
                    or []
                )
                if isinstance(values, (list, tuple)):
                    return [str(x) for x in values]
        except Exception:
            continue

    return []


def safe_strengths(ats_analysis):
    if not identify_strengths:
        return []

    candidates = [
        (ats_analysis,),
        (ats_analysis.get("score_result"),),
    ]

    for args in candidates:
        try:
            result = identify_strengths(*args)

            if isinstance(result, str):
                return [result]

            if isinstance(result, (list, tuple)):
                return [str(x) for x in result]

            if isinstance(result, dict):
                values = (
                    result.get("strengths")
                    or result.get("items")
                    or []
                )
                if isinstance(values, (list, tuple)):
                    return [str(x) for x in values]
        except Exception:
            continue

    return []


def extract_optimized_score(optimized_analysis):
    """
    Get optimized ATS score from the result generated by the scorer.
    """
    if optimized_analysis is None:
        return None

    if isinstance(optimized_analysis, dict):
        # Directly stored by our application.
        if "ats_score" in optimized_analysis:
            return extract_numeric_score(
                optimized_analysis["ats_score"]
            )

        if "score_result" in optimized_analysis:
            result = extract_numeric_score(
                optimized_analysis["score_result"]
            )
            if result is not None:
                return result

    return extract_numeric_score(optimized_analysis)


def get_ai_provider_options():
    """
    Provider names available in the existing environment.
    """
    providers = []

    if get_available_providers:
        try:
            result = get_available_providers()

            if isinstance(result, dict):
                providers = [
                    str(k)
                    for k, enabled in result.items()
                    if enabled
                ]
            elif isinstance(result, (list, tuple, set)):
                providers = [str(x) for x in result]
        except Exception:
            providers = []

    # Fallback to environment variables.
    if not providers:
        if os.getenv("GEMINI_API_KEY"):
            providers.append("Gemini")

        if os.getenv("OPENAI_API_KEY"):
            providers.append("OpenAI")

    # Stable UI order.
    ordered = []
    for provider in ("Gemini", "OpenAI"):
        if provider in providers:
            ordered.append(provider)

    for provider in providers:
        if provider not in ordered:
            ordered.append(provider)

    return ordered


# ============================================================
# STREAMLIT CONFIG
# ============================================================

st.set_page_config(
    page_title="AI Resume Analyzer",
    page_icon="🤖",
    layout="wide",
)

st.title("🤖 AI Resume Analyzer")
st.caption(
    "ATS scoring • Job matching • AI resume optimization • "
    "Same-structure resume editing"
)

# ============================================================
# SESSION STATE
# ============================================================

defaults = {
    "resume_text": "",
    "job_description": "",
    "ats_analysis": None,
    "recommendations": [],
    "strengths": [],
    "uploaded_filename": "",
    "uploaded_bytes": None,
    "uploaded_extension": "",
    "optimized_result": None,
    "optimized_file_bytes": None,
    "optimized_filename": "",
    "optimized_ats": None,
    "optimized_ats_analysis": None,
    "optimized_error": None,
    "last_optimized_source": "",
}

for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value


# ============================================================
# SIDEBAR
# ============================================================

with st.sidebar:
    st.header("⚙️ AI Configuration")

    providers = get_ai_provider_options()

    if providers:
        st.success("AI provider(s) configured")
    else:
        st.error("No AI provider configured")

    selected_provider = None

    if providers:
        selected_provider = st.radio(
            "Choose AI provider",
            providers,
            horizontal=False,
            key="global_provider",
        )

    st.divider()

    st.markdown("### Supported Files")
    st.write("📄 PDF")
    st.write("📝 DOCX")

    st.divider()
    st.caption("Your original resume is never overwritten.")


# ============================================================
# UPLOAD + JD
# ============================================================

st.header("1️⃣ Upload Resume & Job Description")

left, right = st.columns([1, 1])

with left:
    st.subheader("📄 Resume")

    uploaded_file = st.file_uploader(
        "Upload your resume",
        type=["pdf", "docx"],
        key="resume_uploader",
    )

    if uploaded_file is not None:
        suffix = Path(uploaded_file.name).suffix.lower()

        # Reset downstream results when a different file is selected.
        file_token = (
            f"{uploaded_file.name}:"
            f"{len(uploaded_file.getvalue())}"
        )

        if (
            st.session_state.uploaded_filename
            != file_token
        ):
            try:
                resume_text = read_uploaded_resume(uploaded_file)

                st.session_state.resume_text = resume_text
                st.session_state.uploaded_filename = file_token
                st.session_state.uploaded_bytes = (
                    uploaded_file.getvalue()
                )
                st.session_state.uploaded_extension = suffix

                st.session_state.ats_analysis = None
                st.session_state.recommendations = []
                st.session_state.strengths = []
                st.session_state.optimized_result = None
                st.session_state.optimized_file_bytes = None
                st.session_state.optimized_filename = ""
                st.session_state.optimized_ats = None
                st.session_state.optimized_ats_analysis = None
                st.session_state.optimized_error = None

            except Exception as exc:
                st.error(
                    f"Could not read the uploaded resume: {exc}"
                )

        if suffix == ".docx":
            st.success(
                "DOCX detected — structure-preserving "
                "editing available."
            )
        else:
            st.info(
                "PDF detected — ATS analysis is available. "
                "Same-structure editing requires DOCX."
            )

with right:
    st.subheader("📋 Job Description")

    current_jd = st.text_area(
        "Paste the complete job description",
        value=st.session_state.job_description,
        height=330,
        key="jd_input",
    )

    st.session_state.job_description = current_jd


# ============================================================
# ANALYZE
# ============================================================

st.divider()

analyze_clicked = st.button(
    "🔍 Analyze Resume",
    type="primary",
    use_container_width=True,
)

if analyze_clicked:
    if not st.session_state.resume_text.strip():
        st.error("Please upload a valid resume first.")
        st.stop()

    if not st.session_state.job_description.strip():
        st.error("Please paste the complete job description.")
        st.stop()

    with st.spinner("Analyzing resume against the job description..."):
        try:
            ats_analysis = run_ats_analysis(
                st.session_state.resume_text,
                st.session_state.job_description,
            )

            st.session_state.ats_analysis = ats_analysis
            st.session_state.recommendations = (
                safe_recommendations(ats_analysis)
            )
            st.session_state.strengths = (
                safe_strengths(ats_analysis)
            )

            # A new analysis invalidates previous optimized output.
            st.session_state.optimized_result = None
            st.session_state.optimized_file_bytes = None
            st.session_state.optimized_filename = ""
            st.session_state.optimized_ats = None
            st.session_state.optimized_ats_analysis = None
            st.session_state.optimized_error = None

        except Exception as exc:
            st.error(
                "ATS analysis failed."
            )
            with st.expander("Technical error details"):
                st.exception(exc)


# ============================================================
# ATS ANALYSIS DISPLAY
# ============================================================

ats_analysis = st.session_state.ats_analysis

if ats_analysis:
    st.header("2️⃣ ATS Analysis")

    original_score = float(
        ats_analysis.get("ats_score", 0)
    )

    resume_skills = set(
        str(x).strip()
        for x in ats_analysis.get("resume_skills", [])
        if str(x).strip()
    )

    jd_skills = set(
        str(x).strip()
        for x in ats_analysis.get("jd_skills", [])
        if str(x).strip()
    )

    normalized_resume_skills = {
        x.lower()
        for x in resume_skills
    }

    normalized_jd_skills = {
        x.lower()
        for x in jd_skills
    }

    matched_skills = sorted(
        resume_skills
        & {
            jd_skill
            for jd_skill in jd_skills
            if jd_skill.lower() in normalized_resume_skills
        }
    )

    missing_skills = sorted(
        {
            jd_skill
            for jd_skill in jd_skills
            if jd_skill.lower()
            not in normalized_resume_skills
        }
    )

    m1, m2, m3 = st.columns(3)

    with m1:
        st.metric(
            "🎯 ATS Score",
            f"{original_score:.2f}%",
        )

    with m2:
        st.metric(
            "📌 Resume Skills",
            len(resume_skills),
        )

    with m3:
        st.metric(
            "📋 JD Skills",
            len(jd_skills),
        )

    st.subheader("📊 ATS Score Breakdown")

    # Show existing scorer breakdown where available.
    score_result = ats_analysis.get("score_result")

    if isinstance(score_result, dict):
        breakdown_items = []

        possible_breakdowns = [
            ("skills", "Skills"),
            ("skill_score", "Skills"),
            ("semantic", "Semantic"),
            ("semantic_score", "Semantic"),
            ("experience", "Experience"),
            ("experience_score", "Experience"),
            ("keywords", "Keywords"),
            ("keyword_score", "Keywords"),
            ("structure", "Structure"),
            ("structure_score", "Structure"),
        ]

        seen = set()

        for key, label in possible_breakdowns:
            if key not in score_result:
                continue

            if label in seen:
                continue

            value = extract_numeric_score(
                score_result[key]
            )

            if value is None:
                continue

            breakdown_items.append(
                (label, value)
            )
            seen.add(label)

        if breakdown_items:
            columns = st.columns(
                len(breakdown_items)
            )

            for col, (label, value) in zip(
                columns,
                breakdown_items,
            ):
                with col:
                    st.metric(
                        label,
                        f"{value:.2f}%",
                    )
        else:
            st.info(
                "The current scorer did not expose "
                "component-level breakdown values."
            )
    else:
        st.info(
            "The current scorer did not expose "
            "component-level breakdown values."
        )

    st.subheader("✅ Skills Found in Resume")

    if resume_skills:
        cols = st.columns(3)
        for index, skill in enumerate(
            sorted(resume_skills)
        ):
            with cols[index % 3]:
                st.write(f"✓ {skill}")
    else:
        st.info("No resume skills detected.")

    st.subheader("⚠️ Skills Missing from Resume")

    if missing_skills:
        cols = st.columns(3)
        for index, skill in enumerate(
            missing_skills
        ):
            with cols[index % 3]:
                st.write(f"⚠️ {skill}")
    else:
        st.success(
            "No detected JD skills are missing."
        )

    # ============================================================
    # JOB DESCRIPTION ANALYSIS — USER-FRIENDLY VIEW
    # ============================================================

    jd_view = ats_analysis.get(
        "jd_analysis",
        {},
    ) or {}

    st.subheader("📋 Job Requirements")
    st.caption(
        "A clear summary of the employer's requirements and how they "
        "compare with your current resume."
    )

    # ------------------------------------------------------------
    # Top requirement cards
    # ------------------------------------------------------------

    job_title = (
        str(jd_view.get("job_title") or "").strip()
        or "Not detected"
    )

    # Prefer the explicit JD text for ranges such as "2 to 4 years".
    # The displayed range is preserved, while the LOWER bound is used
    # as the minimum experience for eligibility.
    jd_source_text = str(
        st.session_state.get(
            "job_description",
            "",
        ) or ""
    )

    experience_label = "Not specified"
    experience_min_from_range = None
    experience_max_from_range = None

    # Detect common experience-range formats:
    # 2 to 4 years
    # 2-4 years
    # 2–4 years
    # 2—4 years
    experience_match = re.search(
        r"\b(\d+)\s*(?:to|-|\u2013|\u2014)\s*(\d+)\s*years?\b",
        jd_source_text,
        flags=re.IGNORECASE,
    )

    if experience_match:
        experience_min_from_range = float(
            experience_match.group(1)
        )
        experience_max_from_range = float(
            experience_match.group(2)
        )

        experience_label = (
            f"{int(experience_min_from_range)}\u2013"
            f"{int(experience_max_from_range)} years"
        )
    else:
        experience_years = jd_view.get(
            "experience_years"
        )

        if experience_years not in (
            None,
            "",
            0,
        ):
            try:
                experience_value = float(
                    experience_years
                )
            except (TypeError, ValueError):
                experience_value = 0.0

            if experience_value > 0:
                experience_label = (
                    f"{int(experience_value) if experience_value.is_integer() else experience_value:g}+ years"
                )

    # ------------------------------------------------------------
    # EXPERIENCE ELIGIBILITY
    # ------------------------------------------------------------

    # Use the lower bound of an explicit JD range as the minimum
    # requirement. This preserves correct logic for "2 to 4 years":
    # 0-1 years -> Not Eligible
    # 2+ years  -> Eligible
    if experience_min_from_range is not None:
        minimum_experience = experience_min_from_range
    else:
        minimum_experience = jd_view.get(
            "experience_years",
            0,
        )

        try:
            minimum_experience = float(
                minimum_experience or 0
            )
        except (TypeError, ValueError):
            minimum_experience = 0.0

    candidate_experience = 0.0

    try:
        from modules.ats_scorer import (
            _extract_resume_experience_years,
        )

        candidate_experience = (
            _extract_resume_experience_years(
                st.session_state.get(
                    "resume_text",
                    "",
                )
            )
        )

        candidate_experience = float(
            candidate_experience or 0
        )
    except (ImportError, TypeError, ValueError):
        candidate_experience = 0.0
    except Exception:
        # Keep the UI stable if the resume parser cannot determine
        # experience from an unusual resume format.
        candidate_experience = 0.0

    if minimum_experience > 0:
        if candidate_experience >= minimum_experience:
            experience_eligibility = "Eligible"
            experience_eligibility_icon = "✓"
        else:
            experience_eligibility = "Not Eligible"
            experience_eligibility_icon = "✗"
    else:
        experience_eligibility = "Not specified"
        experience_eligibility_icon = "i"

    education_values = jd_view.get(
        "education",
        [],
    )

    education_label = "Not specified"

    if isinstance(
        education_values,
        (list, tuple),
    ):
        education_clean = []
        for value in education_values:
            value = str(value).strip()
            if not value:
                continue

            if value.lower() not in {
                item.lower()
                for item in education_clean
            }:
                education_clean.append(value)

        if education_clean:
            # Collapse common duplicate tokens such as bachelor / bachelor's / degree
            if any(
                "bachelor" in item.lower()
                for item in education_clean
            ):
                education_label = "Bachelor's degree"
            else:
                education_label = " / ".join(
                    education_clean[:3]
                ).title()

    card1, card2, card3 = st.columns(3)

    with card1:
        st.markdown(
            f"""
            <div style="
                border:1px solid #e4e7ec;
                border-radius:14px;
                padding:18px;
                min-height:105px;
                background:#ffffff;
            ">
                <div style="
                    color:#667085;
                    font-size:13px;
                    font-weight:700;
                    text-transform:uppercase;
                    letter-spacing:.04em;
                ">Job Title</div>
                <div style="
                    color:#101828;
                    font-size:19px;
                    font-weight:700;
                    margin-top:7px;
                ">{job_title}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with card2:
        st.markdown(
            textwrap.dedent(
                f"""
                <div style="
                    border:1px solid #e4e7ec;
                    border-radius:14px;
                    padding:18px;
                    min-height:280px;
                    background:#ffffff;
                ">
            <div style="
                color:#667085;
                font-size:13px;
                font-weight:700;
                text-transform:uppercase;
                letter-spacing:.04em;
            ">Experience Required</div>

            <div style="
                color:#101828;
                font-size:19px;
                font-weight:700;
                margin-top:7px;
            ">{experience_label}</div>

            <div style="
                color:#667085;
                font-size:12px;
                font-weight:600;
                margin-top:16px;
            ">MINIMUM REQUIRED</div>

            <div style="
                color:#101828;
                font-size:15px;
                font-weight:700;
                margin-top:4px;
            ">{minimum_experience:g} years</div>

            <div style="
                color:#667085;
                font-size:12px;
                font-weight:600;
                margin-top:16px;
            ">YOUR EXPERIENCE</div>

            <div style="
                color:#101828;
                font-size:15px;
                font-weight:700;
                margin-top:4px;
            ">{candidate_experience:g} years</div>

            <div style="
                color:#667085;
                font-size:12px;
                font-weight:600;
                margin-top:16px;
            ">ELIGIBILITY</div>

            <div style="
                color:#101828;
                font-size:15px;
                font-weight:700;
                margin-top:4px;
            ">
                {experience_eligibility_icon} {experience_eligibility}
            </div>
                </div>
                """
            ),
            unsafe_allow_html=True,
        )

    with card3:
        st.markdown(
            f"""
            <div style="
                border:1px solid #e4e7ec;
                border-radius:14px;
                padding:18px;
                min-height:105px;
                background:#ffffff;
            ">
                <div style="
                    color:#667085;
                    font-size:13px;
                    font-weight:700;
                    text-transform:uppercase;
                    letter-spacing:.04em;
                ">Education</div>
                <div style="
                    color:#101828;
                    font-size:19px;
                    font-weight:700;
                    margin-top:7px;
                ">{education_label}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("#### 🛠 Required Skills")

    skill_categories = jd_view.get(
        "skill_categories",
        {},
    ) or {}

    category_order = [
        ("programming", "Programming"),
        ("web", "Web"),
        ("cloud", "Cloud"),
        ("devops", "DevOps"),
        ("software", "Software"),
        ("business", "Professional Skills"),
        ("other", "Other"),
        ("database", "Database"),
        ("data", "Data"),
        ("ai_ml", "AI / ML"),
    ]

    visible_categories = []

    for key, label in category_order:
        values = skill_categories.get(
            key,
            [],
        )
        if values:
            visible_categories.append(
                (label, values)
            )

    if visible_categories:
        category_columns = st.columns(2)

        for index, (
            label,
            values,
        ) in enumerate(visible_categories):
            with category_columns[
                index % 2
            ]:
                skill_html = "".join(
                    f"""
                    <span style="
                        display:inline-block;
                        padding:6px 10px;
                        margin:3px 4px 3px 0;
                        border-radius:999px;
                        border:1px solid #d0d5dd;
                        background:#f8fafc;
                        color:#344054;
                        font-size:13px;
                    ">{str(value)}</span>
                    """
                    for value in values
                )

                st.markdown(
                    f"""
                    <div style="
                        border:1px solid #eaecf0;
                        border-radius:12px;
                        padding:14px 16px;
                        margin-bottom:10px;
                        background:#fcfcfd;
                    ">
                        <div style="
                            color:#667085;
                            font-size:12px;
                            font-weight:700;
                            text-transform:uppercase;
                            letter-spacing:.04em;
                            margin-bottom:7px;
                        ">{label}</div>
                        <div>{skill_html}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
    else:
        st.info(
            "No categorized skills were detected."
        )

    # ------------------------------------------------------------
    # Resume vs JD skill match
    # ------------------------------------------------------------

    st.markdown("#### 🎯 Your Skill Match")

    resume_skill_map = {
        str(skill).strip().lower():
        str(skill).strip()
        for skill in resume_skills
        if str(skill).strip()
    }

    jd_skill_map = {
        str(skill).strip().lower():
        str(skill).strip()
        for skill in jd_skills
        if str(skill).strip()
    }

    matched_for_display = sorted(
        resume_skill_map[key]
        for key in jd_skill_map
        if key in resume_skill_map
    )

    missing_for_display = sorted(
        jd_skill_map[key]
        for key in jd_skill_map
        if key not in resume_skill_map
    )

    match_col, missing_col = st.columns(2)

    with match_col:
        st.markdown(
            f"""
            <div style="
                border:1px solid #abefc6;
                border-radius:12px;
                padding:15px 16px;
                background:#f6fef9;
                min-height:125px;
            ">
                <div style="
                    color:#067647;
                    font-size:13px;
                    font-weight:700;
                    margin-bottom:8px;
                ">✅ SKILLS FOUND IN YOUR RESUME ({len(matched_for_display)})</div>
                {"".join(
                    f'<span style="display:inline-block;padding:6px 10px;margin:3px 4px 3px 0;border-radius:999px;border:1px solid #abefc6;background:#ecfdf3;color:#067647;font-size:13px;">{skill}</span>'
                    for skill in matched_for_display
                ) or '<span style="color:#667085;">No direct matches detected.</span>'}
            </div>
            """,
            unsafe_allow_html=True,
        )

    with missing_col:
        st.markdown(
            f"""
            <div style="
                border:1px solid #fecdca;
                border-radius:12px;
                padding:15px 16px;
                background:#fffbfa;
                min-height:125px;
            ">
                <div style="
                    color:#b42318;
                    font-size:13px;
                    font-weight:700;
                    margin-bottom:8px;
                ">⚠️ SKILLS NOT FOUND ({len(missing_for_display)})</div>
                {"".join(
                    f'<span style="display:inline-block;padding:6px 10px;margin:3px 4px 3px 0;border-radius:999px;border:1px solid #fecdca;background:#fef3f2;color:#b42318;font-size:13px;">{skill}</span>'
                    for skill in missing_for_display
                ) or '<span style="color:#067647;">All detected JD skills are present.</span>'}
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.caption(
        "Missing skills should only be added to your resume when "
        "you genuinely have experience with them."
    )

    # ------------------------------------------------------------
    # Key responsibilities
    # ------------------------------------------------------------

    st.markdown("#### 💼 Key Responsibilities")

    responsibilities = jd_view.get(
        "responsibilities",
        [],
    ) or []

    boilerplate = {
        "what we expect of you",
        "basic qualifications:",
    }

    cleaned_responsibilities = []

    for responsibility in responsibilities:
        value = str(
            responsibility
        ).strip()

        if not value:
            continue

        if value.lower() in boilerplate:
            continue

        cleaned_responsibilities.append(
            value
        )

    if cleaned_responsibilities:
        for index, responsibility in enumerate(
            cleaned_responsibilities,
            start=1,
        ):
            st.markdown(
                f"""
                <div style="
                    border-left:4px solid #d0d5dd;
                    padding:10px 14px;
                    margin:7px 0;
                    background:#fcfcfd;
                    border-radius:0 9px 9px 0;
                    color:#344054;
                ">
                    <strong>{index}.</strong>&nbsp; {responsibility}
                </div>
                """,
                unsafe_allow_html=True,
            )
    else:
        st.info(
            "No specific responsibilities were detected."
        )

    # Raw JSON remains available only for debugging/development.
    with st.expander("🔧 Developer details — raw JD analysis"):
        st.json(jd_view)

    st.subheader("💪 Resume Strengths")

    strengths = st.session_state.strengths

    if strengths:
        for strength in strengths:
            st.write(f"✓ {strength}")
    else:
        # Reliable fallback strength.
        score_result_dict = score_result
        if isinstance(score_result_dict, dict):
            structure_value = extract_numeric_score(
                score_result_dict.get("structure")
            )
            if structure_value is None:
                structure_value = extract_numeric_score(
                    score_result_dict.get("structure_score")
                )

            if structure_value is not None:
                st.write(
                    f"✓ ATS Structure: "
                    f"{structure_value:.1f}%"
                )
            else:
                st.write(
                    "✓ Resume successfully parsed and analyzed."
                )
        else:
            st.write(
                "✓ Resume successfully parsed and analyzed."
            )

    st.subheader("💡 ATS Recommendations")

    recommendations = st.session_state.recommendations

    if recommendations:
        for recommendation in recommendations:
            st.write(f"• {recommendation}")
    else:
        if original_score < 50:
            st.write(
                "• The resume has low alignment with the "
                "job description. Prioritize relevant skills, "
                "experience and keywords."
            )

        if missing_skills:
            shown = ", ".join(
                missing_skills[:5]
            )
            st.write(
                "• Consider adding the following skills "
                "only if you genuinely have experience with "
                f"them: {shown}."
            )

        st.write(
            "• Improve semantic alignment by using "
            "job-relevant terminology naturally within "
            "your existing experience."
        )

        st.write(
            "• Review the experience section and make "
            "relevant responsibilities and achievements "
            "more explicit."
        )

# ============================================================
# AI OPTIMIZATION
# ============================================================

if ats_analysis:
    st.divider()
    st.header("3️⃣ AI Resume Optimization")

    st.write(
        "The AI will improve the existing resume content "
        "for the selected job description without inventing "
        "experience, skills, achievements, dates or numbers."
    )

    if not selected_provider:
        st.warning(
            "Configure GEMINI_API_KEY or OPENAI_API_KEY "
            "before generating an optimized resume."
        )
    else:
        st.write(
            f"🤖 Selected AI Provider: **{selected_provider}**"
        )

        extension = st.session_state.uploaded_extension

        if extension == ".docx":
            st.success(
                "📝 DOCX detected. The optimizer will "
                "update the existing resume structure."
            )
        else:
            st.warning(
                "PDF resumes can be analyzed, but same-structure "
                "editing requires a DOCX upload."
            )

        generate_clicked = st.button(
            "✨ Generate Optimized Resume",
            type="primary",
            use_container_width=True,
        )

        if generate_clicked:
            if extension != ".docx":
                st.error(
                    "Please upload the resume as DOCX to "
                    "preserve the same structure/layout."
                )
                st.stop()

            if not st.session_state.resume_text.strip():
                st.error("Resume text is empty.")
                st.stop()

            if not st.session_state.job_description.strip():
                st.error("Job description is empty.")
                st.stop()

            with st.spinner(
                "Generating, validating and updating your resume..."
            ):
                try:
                    # 1. Gemini/OpenAI generates optimized content.
                    ai_resume = generate_verified_optimized_resume(
                        st.session_state.resume_text,
                        st.session_state.job_description,
                        ats_analysis,
                        selected_provider,
                    )

                    st.session_state.optimized_result = ai_resume
                    st.session_state.optimized_error = None

                    # 2. Keep original DOCX untouched.
                    original_bytes = (
                        st.session_state.uploaded_bytes
                    )

                    original_path = _write_temp_file(
                        original_bytes,
                        ".docx",
                    )

                    optimized_path = _write_temp_file(
                        b"",
                        ".docx",
                    )

                    try:
                        # 3. Update the EXISTING DOCX structure.
                        update_existing_resume(
                            original_path,
                            optimized_path,
                            ai_resume,
                        )

                        # 4. Read the final generated DOCX.
                        optimized_text = extract_docx_text(
                            optimized_path
                        )

                        if not optimized_text.strip():
                            raise ValueError(
                                "Generated DOCX contains no readable text."
                            )

                        # 5. Re-run the ACTUAL ATS pipeline on
                        #    the generated document.
                        optimized_ats_analysis = run_ats_analysis(
                            optimized_text,
                            st.session_state.job_description,
                        )

                        optimized_score = (
                            optimized_ats_analysis["ats_score"]
                        )

                        # 6. Store the generated file in memory
                        #    for download.
                        with open(
                            optimized_path,
                            "rb",
                        ) as generated_file:
                            generated_bytes = (
                                generated_file.read()
                            )

                        original_name = Path(
                            uploaded_file.name
                        ).stem

                        optimized_filename = (
                            f"{original_name}_"
                            "ATS_Optimized.docx"
                        )

                        st.session_state.optimized_file_bytes = (
                            generated_bytes
                        )

                        st.session_state.optimized_filename = (
                            optimized_filename
                        )

                        st.session_state.optimized_ats = (
                            optimized_score
                        )

                        st.session_state.optimized_ats_analysis = (
                            optimized_ats_analysis
                        )

                        st.session_state.last_optimized_source = (
                            optimized_text
                        )

                    finally:
                        for path in (
                            original_path,
                            optimized_path,
                        ):
                            try:
                                os.unlink(path)
                            except OSError:
                                pass

                    st.success(
                        f"Generated using {selected_provider}. "
                        "Factual validation and ATS re-scoring completed."
                    )

                except Exception as exc:
                    st.session_state.optimized_error = str(
                        exc
                    )

                    st.error(
                        "Resume generation failed."
                    )

                    with st.expander(
                        "Technical error details"
                    ):
                        st.exception(exc)


# ============================================================
# OPTIMIZED RESULT
# ============================================================

if st.session_state.optimized_result:
    st.header("✅ Optimized Resume")

    st.success(
        f"Generated using "
        f"{st.session_state.optimized_result.get('ai_provider', selected_provider)}"
    )

    # ====================================================
    # VALIDATION
    # ====================================================

    validation = st.session_state.get(
        "optimized_result",
        {},
    )

    if not isinstance(validation, dict):
        validation = {}
    else:
        validation = validation.get(
            "validation",
            {},
        )

    if not isinstance(validation, dict):
        validation = {}

    validation_passed = validation.get(
        "passed",
        validation.get(
            "valid",
            True,
        ),
    )

    if validation_passed:
        st.success(
            "🛡️ Factual validation passed. "
            "No unsupported generated claims were detected."
        )
    else:
        st.error(
            "⚠️ Factual validation did not pass."
        )

        with st.expander(
            "🔍 Factual validation details",
            expanded=True,
        ):
            st.write(
                "Unsupported skills:",
                validation.get(
                    "unsupported_skills",
                    [],
                ),
            )

            st.write(
                "Unsupported numbers:",
                validation.get(
                    "unsupported_numbers",
                    [],
                ),
            )

            st.write(
                "Unsupported years:",
                validation.get(
                    "unsupported_years",
                    [],
                ),
            )

            st.write(
                "Unsupported dates:",
                validation.get(
                    "unsupported_dates",
                    [],
                ),
            )

            st.write(
                "Unsupported companies:",
                validation.get(
                    "unsupported_companies",
                    [],
                ),
            )

            st.write(
                "New named terms:",
                validation.get(
                    "new_named_terms",
                    [],
                ),
            )

            st.write(
                "Text overlap:",
                validation.get(
                    "text_overlap",
                    0,
                ),
            )

            st.write(
                "Warnings:",
                validation.get(
                    "warnings",
                    [],
                ),
            )

    # Original ATS analysis is guaranteed to be a dict after a successful
    # analysis, but protect the UI from stale/empty session state.
    if not isinstance(ats_analysis, dict):
        ats_analysis = {}

    original_score = get_score_from_analysis(
        ats_analysis
    )

    optimized_score = (
        st.session_state.optimized_ats
    )

    col1, col2 = st.columns(2)

    with col1:
        st.metric(
            "Original ATS Score",
            f"{original_score:.2f}%",
        )

    with col2:
        if optimized_score is not None:
            delta = (
                float(optimized_score)
                - original_score
            )

            st.metric(
                "Optimized ATS Score",
                f"{float(optimized_score):.2f}%",
                delta=f"{delta:+.2f} points",
            )
        else:
            st.metric(
                "Optimized ATS Score",
                "Not available",
            )

    # --------------------------------------------------------
    # BEFORE vs AFTER
    # --------------------------------------------------------

    optimized_analysis = (
        st.session_state.optimized_ats_analysis
    )

    if optimized_analysis:
        original_skills = {
            str(x).lower(): str(x)
            for x in ats_analysis.get(
                "resume_skills",
                [],
            )
        }

        optimized_skills = {
            str(x).lower(): str(x)
            for x in optimized_analysis.get(
                "resume_skills",
                [],
            )
        }

        jd_skills_after = {
            str(x).lower(): str(x)
            for x in optimized_analysis.get(
                "jd_skills",
                [],
            )
        }

        matched_after = sorted(
            optimized_skills[key]
            for key in jd_skills_after
            if key in optimized_skills
        )

        missing_after = sorted(
            jd_skills_after[key]
            for key in jd_skills_after
            if key not in optimized_skills
        )

        added_skills = sorted(
            optimized_skills[key]
            for key in optimized_skills
            if key not in original_skills
        )

        st.subheader("📊 Before vs After")

        c1, c2, c3, c4 = st.columns(4)

        with c1:
            st.metric(
                "Resume Skills",
                len(optimized_skills),
                delta=(
                    len(optimized_skills)
                    - len(original_skills)
                ),
            )

        with c2:
            st.metric(
                "JD Skills Matched",
                len(matched_after),
            )

        with c3:
            st.metric(
                "Missing JD Skills",
                len(missing_after),
            )

        with c4:
            if optimized_score is not None:
                st.metric(
                    "ATS Improvement",
                    f"{float(optimized_score) - original_score:+.2f}",
                )
            else:
                st.metric(
                    "ATS Improvement",
                    "N/A",
                )

        if added_skills:
            st.subheader("➕ Skills Added to the Resume")
            st.write(
                ", ".join(added_skills)
            )

        if matched_after:
            st.subheader("✅ JD Skills Matched After Optimization")
            st.write(
                ", ".join(matched_after)
            )

        if missing_after:
            st.subheader("⚠️ JD Skills Still Missing")
            st.write(
                ", ".join(missing_after)
            )
        else:
            st.success(
                "All detected JD skills are now present "
                "in the optimized resume."
            )

    # --------------------------------------------------------
    # OPTIMIZED CONTENT PREVIEW
    # --------------------------------------------------------

    st.subheader("👀 Preview Optimized Content")

    ai_result = st.session_state.optimized_result

    preview_sections = [
        ("Professional Summary", "summary"),
        ("Skills", "skills"),
        ("Experience", "experience"),
        ("Projects", "projects"),
        ("Education", "education"),
        ("Certifications", "certifications"),
        ("Achievements", "achievements"),
    ]

    for title, key in preview_sections:
        value = ai_result.get(key)

        if not value:
            continue

        st.markdown(f"**{title}**")

        if isinstance(value, (list, tuple)):
            for item in value:
                st.write(f"• {item}")
        else:
            st.write(str(value))

    # --------------------------------------------------------
    # DOWNLOAD
    # --------------------------------------------------------

    st.subheader("📥 Download")

    if st.session_state.optimized_file_bytes:
        st.download_button(
            label="📄 Download ATS-Optimized Resume",
            data=st.session_state.optimized_file_bytes,
            file_name=(
                st.session_state.optimized_filename
                or "ATS_Optimized_Resume.docx"
            ),
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True,
        )

        st.caption(
            "The original uploaded resume remains unchanged."
        )

    else:
        st.warning(
            "The optimized DOCX is not currently available."
        )


# ============================================================
# ERROR STATE
# ============================================================

if (
    st.session_state.optimized_error
    and not st.session_state.optimized_result
):
    st.error(
        st.session_state.optimized_error
    )
st.markdown(
    """
    <div style="
        text-align: center;
        margin-top: 40px;
        padding: 15px 0;
        color: #777;
        font-size: 14px;
        border-top: 1px solid #ddd;
    ">
        Developed by <strong>Sai Rajan Nandhala</strong>
    </div>
    """,
    unsafe_allow_html=True,
)