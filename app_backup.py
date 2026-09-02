import inspect
import tempfile
from pathlib import Path

import streamlit as st
from docx import Document

from modules.pdf_parser import extract_text_from_pdf
from modules.text_processor import clean_text
from modules.skill_extractor import extract_skills
from modules.jd_analyzer import analyze_job_description
from modules.ats_scorer import calculate_ats_score
from modules.recommendations import (
    generate_recommendations,
    get_resume_strengths,
)
from modules.resume_optimizer import (
    optimize_resume,
    generate_verified_optimized_resume,
)
from modules.resume_editor import update_existing_resume
from modules.ai_rewriter import get_available_providers


# ============================================================
# PAGE CONFIGURATION
# ============================================================

st.set_page_config(
    page_title="AI Resume Analyzer",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ============================================================
# CUSTOM CSS - KEEPING THE ORIGINAL SIMPLE UI
# ============================================================

st.markdown(
    """
    <style>
    .main-title {
        font-size: 42px;
        font-weight: 700;
        margin-bottom: 5px;
    }

    .subtitle {
        font-size: 18px;
        color: #666;
        margin-bottom: 25px;
    }

    .score-card {
        padding: 20px;
        border-radius: 12px;
        border: 1px solid #ddd;
        text-align: center;
    }

    .score-number {
        font-size: 36px;
        font-weight: 700;
    }

    .section-title {
        font-size: 25px;
        font-weight: 650;
        margin-top: 15px;
    }

    .apply-box {
        padding: 18px;
        border-radius: 12px;
        border: 1px solid #ddd;
        margin-top: 10px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# SESSION STATE
# ============================================================

DEFAULT_STATE = {
    "resume_text": "",
    "job_description": "",
    "resume_skills": [],
    "job_skills": [],
    "jd_analysis": {},
    "ats_analysis": {},
    "recommendations": [],
    "strengths": [],
    "optimized_resume": {},
    "resume_filename": "",
    "resume_bytes": None,
    "optimized_resume_bytes": None,
    "optimized_ats_analysis": {},
    "ai_provider": "",
    "application_recommendation": {},
}

for key, value in DEFAULT_STATE.items():
    if key not in st.session_state:
        st.session_state[key] = value


# ============================================================
# HELPERS
# ============================================================

def extract_docx_text(file_or_path):
    """Extract paragraph text from DOCX."""
    document = Document(file_or_path)

    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # Also inspect tables because many resumes store content there.
    for table in document.tables:
        for row in table.rows:
            row_parts = []

            for cell in row.cells:
                cell_text = " ".join(
                    p.text.strip()
                    for p in cell.paragraphs
                    if p.text.strip()
                )

                if cell_text:
                    row_parts.append(cell_text)

            if row_parts:
                parts.append(" | ".join(row_parts))

    return "\n".join(parts)


def normalize_text(text):
    if not text:
        return ""

    return clean_text(str(text)).strip()


def safe_float(value, default=0.0):
    try:
        if value is None:
            return default

        if isinstance(value, str):
            value = value.replace("%", "").strip()

        return float(value)

    except Exception:
        return default


def normalize_score_result(result):
    """
    Normalize different possible ATS scorer return formats into a dict.
    """
    if result is None:
        return {}

    if isinstance(result, dict):
        data = dict(result)

    elif isinstance(result, (int, float)):
        data = {
            "ats_score": float(result)
        }

    else:
        data = {}

        for attr in [
            "ats_score",
            "score",
            "skill_score",
            "semantic_score",
            "experience_score",
            "keyword_score",
            "structure_score",
            "matched_skills",
            "missing_skills",
        ]:
            if hasattr(result, attr):
                data[attr] = getattr(result, attr)

    if "ats_score" not in data:
        for key in [
            "score",
            "overall_score",
            "overall",
            "total_score",
        ]:
            if key in data:
                data["ats_score"] = data[key]
                break

    data["ats_score"] = round(
        safe_float(data.get("ats_score", 0)),
        2,
    )

    return data


def call_ats_scorer(
    resume_text,
    job_description,
    resume_skills,
    job_skills,
    experience_years=0,
    jd_analysis=None,
):
    """
    Call the project's ATS scorer while tolerating common signature
    variations.
    """
    jd_analysis = jd_analysis or {}

    try:
        signature = inspect.signature(
            calculate_ats_score
        )

        parameters = signature.parameters

        kwargs = {}

        if "resume_text" in parameters:
            kwargs["resume_text"] = resume_text

        if "job_description" in parameters:
            kwargs["job_description"] = job_description

        if "resume_skills" in parameters:
            kwargs["resume_skills"] = resume_skills

        if "job_skills" in parameters:
            kwargs["job_skills"] = job_skills

        if "experience_years" in parameters:
            kwargs["experience_years"] = experience_years

        if "jd_analysis" in parameters:
            kwargs["jd_analysis"] = jd_analysis

        if kwargs:
            try:
                return normalize_score_result(
                    calculate_ats_score(**kwargs)
                )
            except TypeError:
                pass

    except Exception:
        pass

    # Fallbacks for positional signatures.
    attempts = [
        (
            resume_text,
            job_description,
            resume_skills,
            job_skills,
            experience_years,
        ),
        (
            resume_text,
            job_description,
            resume_skills,
            job_skills,
        ),
        (
            resume_text,
            job_description,
        ),
    ]

    last_error = None

    for args in attempts:
        try:
            return normalize_score_result(
                calculate_ats_score(*args)
            )
        except TypeError as error:
            last_error = error
            continue

    if last_error:
        raise last_error

    return {}


def calculate_analysis(
    resume_text,
    job_description,
):
    """
    Complete ATS analysis.
    """
    resume_text = normalize_text(resume_text)
    job_description = normalize_text(job_description)

    resume_skills = extract_skills(
        resume_text
    )

    jd_analysis = analyze_job_description(
        job_description
    )

    job_skills = jd_analysis.get(
        "skills",
        [],
    )

    experience_years = jd_analysis.get(
        "experience_years",
        0,
    )

    ats_analysis = call_ats_scorer(
        resume_text,
        job_description,
        resume_skills,
        job_skills,
        experience_years,
        jd_analysis,
    )

    recommendations = generate_recommendations(
        ats_analysis,
        resume_skills,
        job_skills,
    )

    strengths = get_resume_strengths(
        ats_analysis
    )

    return {
        "resume_skills": resume_skills,
        "job_skills": job_skills,
        "jd_analysis": jd_analysis,
        "ats_analysis": ats_analysis,
        "recommendations": recommendations,
        "strengths": strengths,
    }


def calculate_application_recommendation(
    ats_analysis,
    resume_skills,
    job_skills,
    jd_analysis,
    resume_text,
):
    """
    Decide whether the candidate should apply.

    This is deliberately based on multiple signals instead of ATS
    percentage alone.
    """
    ats_score = safe_float(
        ats_analysis.get("ats_score", 0)
    )

    resume_skill_set = {
        str(skill).strip().lower()
        for skill in resume_skills
        if str(skill).strip()
    }

    matched = []

    for skill in job_skills:
        if str(skill).strip().lower() in resume_skill_set:
            matched.append(skill)

    if job_skills:
        skill_match = (
            len(matched)
            / len(job_skills)
        ) * 100
    else:
        skill_match = 0

    missing = [
        skill
        for skill in job_skills
        if str(skill).strip().lower()
        not in resume_skill_set
    ]

    required_experience = safe_float(
        jd_analysis.get(
            "experience_years",
            0,
        )
    )

    # Look for actual experience evidence.
    # This does not claim the candidate has the required years;
    # it only determines whether the resume contains experience
    # information that can be evaluated.
    resume_lower = resume_text.lower()

    experience_terms = [
        "experience",
        "professional experience",
        "work experience",
        "years",
        "year",
    ]

    experience_evidence = any(
        term in resume_lower
        for term in experience_terms
    )

    # Weighted application score.
    application_score = (
        ats_score * 0.60
        + skill_match * 0.30
        + (
            100 if experience_evidence else 0
        ) * 0.10
    )

    if (
        ats_score >= 80
        and skill_match >= 65
        and experience_evidence
    ):
        decision = "APPLY"
        level = "strong"
        emoji = "🟢"
        message = (
            "Your resume is strongly aligned with this "
            "job. You should apply using the optimized resume."
        )

    elif (
        ats_score >= 65
        and skill_match >= 45
    ):
        decision = "APPLY AFTER OPTIMIZATION"
        level = "moderate"
        emoji = "🟡"
        message = (
            "You have a reasonable match for this position. "
            "Use the optimized resume and review the remaining "
            "skill gaps before applying."
        )

    else:
        decision = "LOW MATCH"
        level = "low"
        emoji = "🔴"
        message = (
            "Your resume has limited alignment with this "
            "position. Apply only if you have relevant "
            "experience that is not currently represented "
            "in your resume."
        )

    strengths = []

    if ats_score >= 80:
        strengths.append(
            "Strong ATS alignment"
        )
    elif ats_score >= 65:
        strengths.append(
            "Moderate ATS alignment"
        )

    if skill_match >= 70:
        strengths.append(
            "Strong required-skill match"
        )
    elif skill_match >= 50:
        strengths.append(
            "Good required-skill match"
        )

    if experience_evidence:
        strengths.append(
            "Experience information detected"
        )

    return {
        "decision": decision,
        "level": level,
        "emoji": emoji,
        "message": message,
        "ats_score": round(ats_score, 2),
        "skill_match": round(skill_match, 2),
        "application_score": round(
            application_score,
            2,
        ),
        "matched_skills": matched,
        "missing_skills": missing,
        "strengths": strengths,
        "required_experience": required_experience,
    }


def recalculate_generated_resume(
    optimized_bytes,
    job_description,
):
    """
    Extract the ACTUAL generated DOCX and run the same ATS scorer.
    This prevents the displayed optimized score from being based
    only on the AI response.
    """
    with tempfile.NamedTemporaryFile(
        suffix=".docx",
        delete=False,
    ) as temp_file:

        temp_path = Path(
            temp_file.name
        )

        temp_file.write(
            optimized_bytes
        )

    try:
        generated_text = extract_docx_text(
            str(temp_path)
        )

        generated_text = normalize_text(
            generated_text
        )

        generated_skills = extract_skills(
            generated_text
        )

        jd_analysis = analyze_job_description(
            job_description
        )

        job_skills = jd_analysis.get(
            "skills",
            [],
        )

        experience_years = jd_analysis.get(
            "experience_years",
            0,
        )

        score = call_ats_scorer(
            generated_text,
            job_description,
            generated_skills,
            job_skills,
            experience_years,
            jd_analysis,
        )

        return score

    finally:
        try:
            temp_path.unlink(
                missing_ok=True
            )
        except Exception:
            pass


# ============================================================
# HEADER
# ============================================================

st.markdown(
    '<div class="main-title">🤖 AI Resume Analyzer</div>',
    unsafe_allow_html=True,
)

st.markdown(
    '<div class="subtitle">'
    "ATS scoring • Job matching • AI resume optimization • "
    "Same-structure resume editing"
    "</div>",
    unsafe_allow_html=True,
)


# ============================================================
# SIDEBAR
# ============================================================

with st.sidebar:

    st.header("⚙️ AI Configuration")

    available_providers = (
        get_available_providers()
    )

    if available_providers:

        st.success(
            "AI provider(s) configured"
        )

        ai_provider = st.radio(
            "Choose AI provider",
            available_providers,
            index=0,
        )

        st.session_state[
            "ai_provider"
        ] = ai_provider

    else:

        st.warning(
            "No AI provider configured."
        )

        st.info(
            "Add GEMINI_API_KEY or "
            "OPENAI_API_KEY to your .env file."
        )

        ai_provider = None

    st.divider()

    st.subheader("Supported Files")

    st.write("📄 PDF")
    st.write("📝 DOCX")

    st.divider()

    st.caption(
        "Your original resume is never overwritten."
    )


# ============================================================
# INPUT SECTION
# ============================================================

st.header("1️⃣ Upload Resume & Job Description")

resume_col, jd_col = st.columns(
    [1, 1]
)


# ============================================================
# RESUME
# ============================================================

with resume_col:

    st.subheader("📄 Resume")

    resume_file = st.file_uploader(
        "Upload your resume",
        type=["pdf", "docx"],
        help=(
            "DOCX is recommended when you want the "
            "original resume structure preserved."
        ),
    )

    if resume_file:

        file_extension = (
            Path(
                resume_file.name
            ).suffix.lower()
        )

        if file_extension == ".docx":

            st.success(
                "✅ DOCX detected — "
                "structure-preserving editing available."
            )

        elif file_extension == ".pdf":

            st.info(
                "📄 PDF detected — "
                "analysis is available. "
                "Upload DOCX for same-structure editing."
            )


# ============================================================
# JOB DESCRIPTION
# ============================================================

with jd_col:

    st.subheader("📋 Job Description")

    job_description = st.text_area(
        "Paste the complete job description",
        height=260,
        placeholder=(
            "Paste the complete job description here..."
        ),
    )


# ============================================================
# ANALYZE BUTTON
# ============================================================

st.divider()

analyze_button = st.button(
    "🔍 Analyze Resume",
    type="primary",
    use_container_width=True,
)


# ============================================================
# ANALYSIS
# ============================================================

if analyze_button:

    if resume_file is None:

        st.error(
            "❌ Please upload your resume."
        )

        st.stop()

    if not job_description.strip():

        st.error(
            "❌ Please paste the job description."
        )

        st.stop()

    with st.spinner(
        "🤖 Analyzing resume and job description..."
    ):

        try:

            file_extension = (
                Path(
                    resume_file.name
                ).suffix.lower()
            )

            if file_extension == ".pdf":

                raw_resume_text = (
                    extract_text_from_pdf(
                        resume_file
                    )
                )

            elif file_extension == ".docx":

                raw_resume_text = (
                    extract_docx_text(
                        resume_file
                    )
                )

            else:

                raise ValueError(
                    "Unsupported resume format."
                )

            if not raw_resume_text.strip():

                st.error(
                    "❌ Could not extract text from the resume."
                )

                st.stop()

            resume_text = normalize_text(
                raw_resume_text
            )

            analysis = calculate_analysis(
                resume_text,
                job_description,
            )

            st.session_state[
                "resume_text"
            ] = resume_text

            st.session_state[
                "job_description"
            ] = job_description

            st.session_state[
                "resume_skills"
            ] = analysis[
                "resume_skills"
            ]

            st.session_state[
                "job_skills"
            ] = analysis[
                "job_skills"
            ]

            st.session_state[
                "jd_analysis"
            ] = analysis[
                "jd_analysis"
            ]

            st.session_state[
                "ats_analysis"
            ] = analysis[
                "ats_analysis"
            ]

            st.session_state[
                "recommendations"
            ] = analysis[
                "recommendations"
            ]

            st.session_state[
                "strengths"
            ] = analysis[
                "strengths"
            ]

            # Keep the original uploaded file exactly as supplied.
            st.session_state[
                "resume_filename"
            ] = resume_file.name

            st.session_state[
                "resume_bytes"
            ] = resume_file.getvalue()

            st.session_state[
                "optimized_resume"
            ] = {}

            st.session_state[
                "optimized_resume_bytes"
            ] = None

            st.session_state[
                "optimized_ats_analysis"
            ] = {}

            st.session_state[
                "application_recommendation"
            ] = {}

            st.success(
                "✅ Resume analysis completed successfully!"
            )

        except Exception as error:

            st.error(
                f"❌ Analysis failed: {error}"
            )

            with st.expander(
                "Technical error details"
            ):

                st.exception(error)

            st.stop()


# ============================================================
# RESULTS
# ============================================================

if st.session_state[
    "ats_analysis"
]:

    ats = st.session_state[
        "ats_analysis"
    ]

    resume_skills = st.session_state[
        "resume_skills"
    ]

    job_skills = st.session_state[
        "job_skills"
    ]

    recommendations = st.session_state[
        "recommendations"
    ]

    strengths = st.session_state[
        "strengths"
    ]

    optimized_resume = st.session_state[
        "optimized_resume"
    ]


    # ========================================================
    # ATS SCORE
    # ========================================================

    st.divider()

    st.header("2️⃣ ATS Analysis")

    score_col1, score_col2, score_col3, score_col4 = (
        st.columns(4)
    )

    with score_col1:

        st.metric(
            "🎯 ATS Score",
            f"{ats.get('ats_score', 0)}%",
        )

    with score_col2:

        st.metric(
            "📌 Resume Skills",
            len(resume_skills),
        )

    with score_col3:

        st.metric(
            "📋 JD Skills",
            len(job_skills),
        )

    with score_col4:

        st.metric(
            "✅ Matched Skills",
            len(
                ats.get(
                    "matched_skills",
                    [],
                )
            ),
        )


    # ========================================================
    # SCORE BREAKDOWN
    # ========================================================

    st.subheader(
        "📊 ATS Score Breakdown"
    )

    breakdown1, breakdown2, breakdown3, breakdown4, breakdown5 = (
        st.columns(5)
    )

    with breakdown1:

        st.metric(
            "Skills",
            f"{ats.get('skill_score', 0)}%",
        )

    with breakdown2:

        st.metric(
            "Semantic",
            f"{ats.get('semantic_score', 0)}%",
        )

    with breakdown3:

        st.metric(
            "Experience",
            f"{ats.get('experience_score', 0)}%",
        )

    with breakdown4:

        st.metric(
            "Keywords",
            f"{ats.get('keyword_score', 0)}%",
        )

    with breakdown5:

        st.metric(
            "Structure",
            f"{ats.get('structure_score', 0)}%",
        )


    # ========================================================
    # MATCHED / MISSING SKILLS
    # ========================================================

    matched_col, missing_col = st.columns(
        2
    )

    with matched_col:

        st.subheader(
            "✅ Skills Found in Resume"
        )

        matched_skills = ats.get(
            "matched_skills",
            [],
        )

        if matched_skills:

            for skill in matched_skills:

                st.write(
                    f"✓ {skill}"
                )

        else:

            st.info(
                "No matching skills detected."
            )

    with missing_col:

        st.subheader(
            "⚠️ Skills Missing from Resume"
        )

        missing_skills = ats.get(
            "missing_skills",
            [],
        )

        if missing_skills:

            for skill in missing_skills:

                st.write(
                    f"⚠️ {skill}"
                )

        else:

            st.success(
                "No major missing skills detected."
            )


    # ========================================================
    # JD DETAILS
    # ========================================================

    st.divider()

    st.subheader(
        "📋 Job Description Analysis"
    )

    jd_analysis = st.session_state[
        "jd_analysis"
    ]

    with st.expander(
        "View detailed JD analysis"
    ):

        experience_years = (
            jd_analysis.get(
                "experience_years",
                0,
            )
        )

        st.write(
            "**Required Experience:**"
        )

        if experience_years:

            st.write(
                f"{experience_years}+ years"
            )

        else:

            st.write(
                "Not detected"
            )

        st.write(
            "**Education Requirements:**"
        )

        education = jd_analysis.get(
            "education",
            [],
        )

        if education:

            st.write(
                ", ".join(
                    education
                )
            )

        else:

            st.write(
                "Not detected"
            )

        st.write(
            "**Responsibilities:**"
        )

        responsibilities = (
            jd_analysis.get(
                "responsibilities",
                [],
            )
        )

        if responsibilities:

            for responsibility in responsibilities:

                st.write(
                    f"• {responsibility}"
                )

        else:

            st.write(
                "No responsibilities detected."
            )


    # ========================================================
    # STRENGTHS
    # ========================================================

    if strengths:

        st.subheader(
            "💪 Resume Strengths"
        )

        for strength in strengths:

            st.write(
                f"✓ {strength}"
            )


    # ========================================================
    # RECOMMENDATIONS
    # ========================================================

    st.subheader(
        "💡 ATS Recommendations"
    )

    if recommendations:

        for recommendation in recommendations:

            st.write(
                f"• {recommendation}"
            )

    else:

        st.success(
            "Your resume does not have major ATS issues."
        )


    # ========================================================
    # AI OPTIMIZATION
    # ========================================================

    st.divider()

    st.header(
        "3️⃣ AI Resume Optimization"
    )

    st.write(
        "The AI will improve the existing resume content "
        "for the selected job description without changing "
        "your role, section headings, company names, dates, "
        "education, projects, or factual information."
    )

    st.info(
        "🔒 Structure protection: existing headings and role "
        "names are preserved. Only supported content and "
        "wording are optimized."
    )


    # ========================================================
    # PROVIDER
    # ========================================================

    if available_providers:

        selected_provider = st.radio(
            "🤖 AI Provider",
            available_providers,
            index=(
                available_providers.index(
                    st.session_state.get(
                        "ai_provider",
                        available_providers[0],
                    )
                )
                if st.session_state.get(
                    "ai_provider"
                ) in available_providers
                else 0
            ),
            horizontal=True,
        )

        st.session_state[
            "ai_provider"
        ] = selected_provider

    else:

        st.error(
            "No AI provider is configured."
        )

        st.info(
            "Add GEMINI_API_KEY or OPENAI_API_KEY "
            "to your .env file."
        )

        selected_provider = None


    # ========================================================
    # INPUT FORMAT WARNING
    # ========================================================

    resume_filename = st.session_state[
        "resume_filename"
    ]

    if resume_filename.lower().endswith(
        ".docx"
    ):

        st.success(
            "📝 DOCX detected. The optimizer will update "
            "the existing resume structure without replacing "
            "your headings or role names."
        )

    else:

        st.warning(
            "📄 PDF detected. ATS analysis works, but "
            "same-structure editing requires the original DOCX file."
        )


    # ========================================================
    # PREVIEW
    # ========================================================

    with st.expander(
        "👀 Preview optimized content"
    ):

        st.subheader(
            "Professional Summary"
        )

        st.write(
            optimized_resume.get(
                "summary",
                "",
            )
        )

        st.subheader(
            "Skills"
        )

        preview_skills = (
            optimized_resume.get(
                "skills",
                [],
            )
        )

        if preview_skills:

            st.write(
                " • ".join(
                    preview_skills
                )
            )

        preview_experience = (
            optimized_resume.get(
                "experience",
                [],
            )
        )

        if preview_experience:

            st.subheader(
                "Professional Experience"
            )

            for item in preview_experience:

                st.write(
                    f"• {item}"
                )

        preview_projects = (
            optimized_resume.get(
                "projects",
                [],
            )
        )

        if preview_projects:

            st.subheader(
                "Projects"
            )

            for item in preview_projects:

                st.write(
                    f"• {item}"
                )


    # ========================================================
    # GENERATE BUTTON
    # ========================================================

    generate_button = st.button(
        "✨ Generate ATS-Optimized Resume",
        type="primary",
        use_container_width=True,
        disabled=(
            selected_provider is None
        ),
    )


    # ========================================================
    # GENERATION
    # ========================================================

    if generate_button:

        if not resume_filename.lower().endswith(
            ".docx"
        ):

            st.error(
                "❌ Same-structure resume generation "
                "requires a DOCX input."
            )

            st.info(
                "Please upload your resume as DOCX "
                "and run the analysis again."
            )

            st.stop()

        with st.spinner(
            f"🤖 {selected_provider} is optimizing, "
            "validating and re-scoring your resume..."
        ):

            try:

                # ------------------------------------------------
                # AI rewrite + factual validation
                # ------------------------------------------------

                ai_resume = (
                    generate_verified_optimized_resume(
                        st.session_state[
                            "resume_text"
                        ],
                        st.session_state[
                            "job_description"
                        ],
                        st.session_state[
                            "ats_analysis"
                        ],
                        selected_provider,
                    )
                )

                # ------------------------------------------------
                # TEMPORARY WORKSPACE
                # ------------------------------------------------

                with tempfile.TemporaryDirectory() as temp_dir:

                    temp_dir = Path(
                        temp_dir
                    )

                    original_path = (
                        temp_dir
                        / "original_resume.docx"
                    )

                    optimized_path = (
                        temp_dir
                        / "ATS_Optimized_Resume.docx"
                    )

                    # Save uploaded original.
                    original_path.write_bytes(
                        st.session_state[
                            "resume_bytes"
                        ]
                    )

                    # IMPORTANT:
                    # This edits the existing DOCX instead of
                    # creating a new resume from scratch.
                    updated_path = (
                        update_existing_resume(
                            original_path,
                            optimized_path,
                            ai_resume,
                        )
                    )

                    optimized_bytes = (
                        Path(
                            updated_path
                        ).read_bytes()
                    )

                # ------------------------------------------------
                # ACTUAL ATS RE-SCORE
                # ------------------------------------------------

                optimized_ats = (
                    recalculate_generated_resume(
                        optimized_bytes,
                        st.session_state[
                            "job_description"
                        ],
                    )
                )

                # ------------------------------------------------
                # Save result
                # ------------------------------------------------

                st.session_state[
                    "optimized_resume"
                ] = ai_resume

                st.session_state[
                    "optimized_resume_bytes"
                ] = optimized_bytes

                st.session_state[
                    "optimized_ats_analysis"
                ] = optimized_ats

                # ------------------------------------------------
                # Application recommendation
                # ------------------------------------------------

                application_result = (
                    calculate_application_recommendation(
                        optimized_ats,
                        extract_skills(
                            extract_docx_text(
                                tempfile.NamedTemporaryFile(
                                    suffix=".docx",
                                    delete=False,
                                ).name
                            )
                        )
                        if False
                        else st.session_state[
                            "resume_skills"
                        ],
                        st.session_state[
                            "job_skills"
                        ],
                        st.session_state[
                            "jd_analysis"
                        ],
                        st.session_state[
                            "resume_text"
                        ],
                    )
                )

                st.session_state[
                    "application_recommendation"
                ] = application_result

                st.success(
                    "🎉 ATS-optimized resume generated successfully!"
                )

            except Exception as error:

                st.error(
                    "❌ Resume generation failed."
                )

                with st.expander(
                    "Technical error details"
                ):

                    st.exception(error)


    # ========================================================
    # GENERATED RESULT
    # ========================================================

    if st.session_state[
        "optimized_resume_bytes"
    ] is not None:

        st.divider()

        st.header(
            "4️⃣ Optimized Resume"
        )

        provider_used = (
            st.session_state[
                "optimized_resume"
            ].get(
                "ai_provider",
                st.session_state.get(
                    "ai_provider",
                    "AI",
                ),
            )
        )

        st.success(
            f"🤖 Generated using {provider_used}"
        )


        # ====================================================
        # VALIDATION
        # ====================================================

        validation = (
            st.session_state[
                "optimized_resume"
            ].get(
                "validation",
                {},
            )
        )

        if validation.get(
            "valid",
            False,
        ):

            st.success(
                "🛡️ Factual validation passed. "
                "No unsupported generated claims were detected."
            )

        else:

            st.error(
                "⚠️ Factual validation did not pass."
            )


        # ====================================================
        # BEFORE / AFTER ATS
        # ====================================================

        before_score = safe_float(
            ats.get(
                "ats_score",
                0,
            )
        )

        after_score = safe_float(
            st.session_state[
                "optimized_ats_analysis"
            ].get(
                "ats_score",
                0,
            )
        )

        before_col, after_col = st.columns(
            2
        )

        with before_col:

            st.metric(
                "Original ATS Score",
                f"{before_score:.2f}%",
            )

        with after_col:

            st.metric(
                "Optimized ATS Score",
                f"{after_score:.2f}%",
                delta=round(
                    after_score
                    - before_score,
                    2,
                ),
            )


        # ====================================================
        # OPTIMIZATION SUMMARY
        # ====================================================

        st.subheader(
            "📊 Optimization Summary"
        )

        improvement = (
            after_score
            - before_score
        )

        summary_col1, summary_col2, summary_col3 = (
            st.columns(3)
        )

        with summary_col1:

            st.metric(
                "Original",
                f"{before_score:.2f}%",
            )

        with summary_col2:

            st.metric(
                "Optimized",
                f"{after_score:.2f}%",
            )

        with summary_col3:

            st.metric(
                "Improvement",
                f"{improvement:+.2f} points",
            )


        # ====================================================
        # APPLICATION RECOMMENDATION
        # ====================================================

        application_result = (
            st.session_state.get(
                "application_recommendation",
                {},
            )
        )

        # Recalculate using final score if needed.
        if not application_result:

            application_result = (
                calculate_application_recommendation(
                    st.session_state[
                        "optimized_ats_analysis"
                    ],
                    st.session_state[
                        "resume_skills"
                    ],
                    st.session_state[
                        "job_skills"
                    ],
                    st.session_state[
                        "jd_analysis"
                    ],
                    st.session_state[
                        "resume_text"
                    ],
                )
            )

            st.session_state[
                "application_recommendation"
            ] = application_result

        st.divider()

        st.subheader(
            "🎯 Should You Apply?"
        )

        decision = application_result.get(
            "decision",
            "REVIEW",
        )

        level = application_result.get(
            "level",
            "moderate",
        )

        emoji = application_result.get(
            "emoji",
            "🟡",
        )

        if level == "strong":

            st.success(
                f"{emoji} **{decision}**"
            )

        elif level == "moderate":

            st.warning(
                f"{emoji} **{decision}**"
            )

        else:

            st.error(
                f"{emoji} **{decision}**"
            )

        st.write(
            application_result.get(
                "message",
                "",
            )
        )

        apply_col1, apply_col2, apply_col3 = (
            st.columns(3)
        )

        with apply_col1:

            st.metric(
                "Final ATS Score",
                f"{application_result.get('ats_score', after_score):.2f}%",
            )

        with apply_col2:

            st.metric(
                "JD Skill Match",
                f"{application_result.get('skill_match', 0):.2f}%",
            )

        with apply_col3:

            st.metric(
                "Application Score",
                f"{application_result.get('application_score', 0):.2f}%",
            )

        application_strengths = (
            application_result.get(
                "strengths",
                [],
            )
        )

        if application_strengths:

            st.write(
                "**Why:**"
            )

            for item in application_strengths:

                st.write(
                    f"✓ {item}"
                )

        application_gaps = (
            application_result.get(
                "missing_skills",
                [],
            )
        )

        if application_gaps:

            st.write(
                "**⚠️ Important skill gaps:**"
            )

            for item in application_gaps[:10]:

                st.write(
                    f"• {item}"
                )


        # ====================================================
        # DOWNLOAD
        # ====================================================

        st.subheader(
            "📥 Download"
        )

        st.download_button(
            label="📄 Download ATS-Optimized Resume",
            data=st.session_state[
                "optimized_resume_bytes"
            ],
            file_name=(
                "ATS_Optimized_Resume.docx"
            ),
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True,
        )

        st.caption(
            "The original uploaded resume remains unchanged. "
            "Existing section headings and role names are preserved."
        )
# ============================================================
# FOOTER
# ============================================================

st.divider()

st.markdown(
    """
    <div style="text-align: center; padding: 20px 0 10px 0;">
        <p style="font-size: 15px; color: #666;">
            Developed by <strong>Sai Rajan Nandhala</strong>
        </p>
        <p style="font-size: 12px; color: #999;">
            AI Resume Analyzer • ATS Optimization • Job Matching
        </p>
    </div>
    """,
    unsafe_allow_html=True,
)