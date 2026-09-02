from copy import deepcopy
from pathlib import Path

from docx import Document


SECTION_HEADINGS = {
    "summary": {
        "summary",
        "professional summary",
        "profile",
        "objective",
    },
    "skills": {
        "skills",
        "technical skills",
        "core skills",
        "core competencies",
        "key skills",
    },
    "experience": {
        "experience",
        "work experience",
        "professional experience",
        "employment",
    },
    "projects": {
        "projects",
        "academic projects",
        "personal projects",
    },
    "education": {
        "education",
        "academic background",
        "qualifications",
    },
    "certifications": {
        "certifications",
        "certificates",
    },
    "achievements": {
        "achievements",
        "accomplishments",
    },
}


def normalize_heading(text):
    """Normalize section heading text."""

    return (
        text
        .strip()
        .lower()
        .rstrip(":")
    )


def detect_section_heading(text):
    """Return the internal section name for a heading."""

    normalized = normalize_heading(text)

    for section_name, headings in SECTION_HEADINGS.items():

        if normalized in {
            normalize_heading(item)
            for item in headings
        }:
            return section_name

    return None


def load_docx_resume(file_path):
    """Load the original DOCX resume."""

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"Resume file not found: {path}"
        )

    if path.suffix.lower() != ".docx":
        raise ValueError(
            "The structure-preserving editor currently "
            "requires a DOCX file."
        )

    return Document(str(path))


def get_section_ranges(document):
    """
    Identify section boundaries using paragraph indexes.

    The original order is preserved.
    """

    sections = []

    current_section = "header"
    current_heading_index = None
    current_content_indexes = []

    paragraphs = document.paragraphs

    for index, paragraph in enumerate(paragraphs):

        text = paragraph.text.strip()

        if not text:
            continue

        detected = detect_section_heading(text)

        if detected:

            if current_heading_index is not None:

                sections.append({
                    "name": current_section,
                    "heading_index": current_heading_index,
                    "content_indexes": current_content_indexes,
                })

            current_section = detected
            current_heading_index = index
            current_content_indexes = []

        elif current_heading_index is not None:

            current_content_indexes.append(index)

        else:

            current_content_indexes.append(index)

    if current_heading_index is not None:

        sections.append({
            "name": current_section,
            "heading_index": current_heading_index,
            "content_indexes": current_content_indexes,
        })

    return sections


def get_section(
    document,
    section_name
):
    """Find a particular section."""

    sections = get_section_ranges(
        document
    )

    for section in sections:

        if section["name"] == section_name:
            return section

    return None


def copy_paragraph_format(
    source,
    target
):
    """Copy paragraph-level formatting."""

    target.alignment = source.alignment

    target.paragraph_format.left_indent = (
        source.paragraph_format.left_indent
    )

    target.paragraph_format.right_indent = (
        source.paragraph_format.right_indent
    )

    target.paragraph_format.first_line_indent = (
        source.paragraph_format.first_line_indent
    )

    target.paragraph_format.space_before = (
        source.paragraph_format.space_before
    )

    target.paragraph_format.space_after = (
        source.paragraph_format.space_after
    )

    target.paragraph_format.line_spacing = (
        source.paragraph_format.line_spacing
    )


def copy_run_format(
    source_run,
    target_run
):
    """Copy character-level formatting."""

    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline

    target_run.font.name = (
        source_run.font.name
    )

    target_run.font.size = (
        source_run.font.size
    )

    if source_run.font.color.type is not None:

        target_run.font.color.type = (
            source_run.font.color.type
        )

        if source_run.font.color.rgb:

            target_run.font.color.rgb = (
                source_run.font.color.rgb
            )


def replace_paragraph_text(
    paragraph,
    new_text
):
    """
    Replace text while preserving the formatting
    of the first run.
    """

    if not paragraph.runs:

        run = paragraph.add_run(
            new_text
        )

        return

    template_run = paragraph.runs[0]

    for run in paragraph.runs:

        run.text = ""

    template_run.text = new_text


def clone_paragraph(
    document,
    source_paragraph,
    text
):
    """
    Create a new paragraph based on an existing
    paragraph's XML and formatting.
    """

    new_paragraph = document.add_paragraph()

    copy_paragraph_format(
        source_paragraph,
        new_paragraph
    )

    # Copy paragraph style.
    try:
        new_paragraph.style = (
            source_paragraph.style
        )
    except Exception:
        pass

    run = new_paragraph.add_run(
        text
    )

    if source_paragraph.runs:

        copy_run_format(
            source_paragraph.runs[0],
            run
        )

    return new_paragraph


def insert_paragraph_after(
    paragraph,
    text
):
    """
    Insert a paragraph immediately after another
    paragraph while copying its formatting.
    """

    new_p = deepcopy(
        paragraph._p
    )

    paragraph._p.addnext(
        new_p
    )

    from docx.text.paragraph import Paragraph

    new_paragraph = Paragraph(
        new_p,
        paragraph._parent
    )

    replace_paragraph_text(
        new_paragraph,
        text
    )

    return new_paragraph


def is_bullet_paragraph(paragraph):
    """
    Detect whether a paragraph is using a bullet style.
    """

    style_name = ""

    try:
        style_name = (
            paragraph.style.name
            or ""
        ).lower()
    except Exception:
        pass

    if "bullet" in style_name:
        return True

    # Check numbering XML as a fallback.
    p_pr = paragraph._p.pPr

    if p_pr is not None:

        num_pr = p_pr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
        )

        if num_pr is not None:
            return True

    return False


def split_content_lines(content):
    """
    Convert generated content into individual lines.
    """

    if not content:
        return []

    if isinstance(content, str):

        lines = content.splitlines()

    elif isinstance(content, list):

        lines = []

        for item in content:

            if item is None:
                continue

            lines.extend(
                str(item).splitlines()
            )

    else:

        lines = [
            str(content)
        ]

    return [
        line.strip()
        for line in lines
        if line.strip()
    ]


def update_section_content(
    document,
    section_name,
    new_content
):
    """
    Update the content of an existing section.

    Existing paragraphs are reused first.
    Additional paragraphs are inserted when necessary.

    The section heading itself is never replaced.
    """

    section = get_section(
        document,
        section_name
    )

    if not section:
        return False

    content_indexes = section[
        "content_indexes"
    ]

    new_lines = split_content_lines(
        new_content
    )

    if not new_lines:
        return False

    paragraphs = document.paragraphs

    existing_paragraphs = [
        paragraphs[index]
        for index in content_indexes
        if index < len(paragraphs)
    ]

    # -----------------------------------------
    # Reuse existing paragraphs
    # -----------------------------------------

    for position, line in enumerate(
        new_lines
    ):

        if position < len(
            existing_paragraphs
        ):

            paragraph = (
                existing_paragraphs[position]
            )

            replace_paragraph_text(
                paragraph,
                line
            )

    # -----------------------------------------
    # Add additional paragraphs
    # -----------------------------------------

    if len(new_lines) > len(
        existing_paragraphs
    ):

        if existing_paragraphs:

            template = (
                existing_paragraphs[-1]
            )

        else:

            template = paragraphs[
                section["heading_index"]
            ]

        for line in new_lines[
            len(existing_paragraphs):
        ]:

            insert_paragraph_after(
                template,
                line
            )

            # Newly inserted paragraph becomes
            # the template for the next paragraph.
            template = document.paragraphs[
                document.paragraphs.index(template) + 1
            ]

    # -----------------------------------------
    # Remove excess old paragraphs
    # -----------------------------------------

    if len(existing_paragraphs) > len(
        new_lines
    ):

        for paragraph in existing_paragraphs[
            len(new_lines):
        ]:

            # Never remove the section heading.
            parent = paragraph._element.getparent()

            if parent is not None:

                parent.remove(
                    paragraph._element
                )

    return True


def update_skills(
    document,
    skills
):
    """Update the existing skills section."""

    if not skills:
        return False

    text = " • ".join(
        str(skill)
        for skill in skills
    )

    return update_section_content(
        document,
        "skills",
        [text]
    )


def update_summary(
    document,
    summary
):
    """Update the existing summary section."""

    return update_section_content(
        document,
        "summary",
        summary
    )


def update_experience(
    document,
    experience
):
    """Update experience while preserving bullet structure."""

    return update_section_content(
        document,
        "experience",
        experience
    )


def update_projects(
    document,
    projects
):
    """Update project content."""

    return update_section_content(
        document,
        "projects",
        projects
    )


def update_existing_resume(
    input_path,
    output_path,
    optimized_resume
):
    """
    Update the original DOCX resume.

    The original file is NEVER overwritten.
    """

    document = load_docx_resume(
        input_path
    )

    # -----------------------------
    # Summary
    # -----------------------------

    update_summary(
        document,
        optimized_resume.get(
            "summary",
            ""
        )
    )

    # -----------------------------
    # Skills
    # -----------------------------

    update_skills(
        document,
        optimized_resume.get(
            "skills",
            []
        )
    )

    # -----------------------------
    # Experience
    # -----------------------------

    update_experience(
        document,
        optimized_resume.get(
            "experience",
            []
        )
    )

    # -----------------------------
    # Projects
    # -----------------------------

    update_projects(
        document,
        optimized_resume.get(
            "projects",
            []
        )
    )

    # -----------------------------
    # Save to a new file
    # -----------------------------

    output_path = Path(
        output_path
    )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    document.save(
        str(output_path)
    )

    return str(output_path)