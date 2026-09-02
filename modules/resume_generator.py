"""
resume_generator.py

Fallback resume generation utilities for the AI Resume Analyzer.

Important:
The main application should use resume_editor.py for same-structure
DOCX editing. This module is intended for:
- structured resume generation
- text/preview generation
- fallback DOCX generation when a new document is explicitly needed

It does not replace the original resume editor.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# ============================================================
# HELPERS
# ============================================================

def _clean_text(value: Any) -> str:
    """Convert a value to clean text."""

    if value is None:
        return ""

    return str(value).strip()


def _to_lines(value: Any) -> List[str]:
    """
    Convert strings, lists and dictionaries into readable lines.
    """

    if value is None:
        return []

    if isinstance(value, dict):
        name = _clean_text(
            value.get("name")
            or value.get("title")
            or value.get("project_name")
            or ""
        )

        description = _clean_text(
            value.get("description")
            or value.get("details")
            or value.get("content")
            or value.get("text")
            or ""
        )

        lines = []

        if name:
            lines.append(name)

        if description:
            lines.extend(
                line.strip()
                for line in description.splitlines()
                if line.strip()
            )

        if lines:
            return lines

        return [
            str(item).strip()
            for item in value.values()
            if str(item).strip()
        ]

    if isinstance(value, (list, tuple, set)):
        lines = []

        for item in value:
            lines.extend(
                _to_lines(item)
            )

        return lines

    text = _clean_text(value)

    if not text:
        return []

    return [
        line.strip("•- \t")
        for line in text.splitlines()
        if line.strip()
    ]


# ============================================================
# STRUCTURED RESUME
# ============================================================

def normalize_resume(
    resume: Dict[str, Any],
) -> Dict[str, Any]:
    """
    Normalize a structured resume dictionary.
    """

    if not isinstance(
        resume,
        dict,
    ):
        raise ValueError(
            "Resume must be a dictionary."
        )

    skills = resume.get(
        "skills",
        [],
    )

    if isinstance(
        skills,
        str,
    ):
        skills = _to_lines(
            skills
        )

    if not isinstance(
        skills,
        (list, tuple),
    ):
        skills = []

    normalized_skills = []

    seen = set()

    for skill in skills:
        skill_text = _clean_text(
            skill
        )

        if not skill_text:
            continue

        key = skill_text.lower()

        if key not in seen:
            seen.add(key)
            normalized_skills.append(
                skill_text
            )

    return {
        "summary": _clean_text(
            resume.get(
                "summary",
                "",
            )
        ),
        "skills": normalized_skills,
        "experience": _to_lines(
            resume.get(
                "experience",
                [],
            )
        ),
        "projects": _to_lines(
            resume.get(
                "projects",
                [],
            )
        ),
        "education": _to_lines(
            resume.get(
                "education",
                [],
            )
        ),
        "certifications": _to_lines(
            resume.get(
                "certifications",
                [],
            )
        ),
        "achievements": _to_lines(
            resume.get(
                "achievements",
                [],
            )
        ),
    }


# ============================================================
# TEXT GENERATION
# ============================================================

def resume_to_text(
    resume: Dict[str, Any],
) -> str:
    """
    Convert a structured resume into readable text.
    """

    resume = normalize_resume(
        resume
    )

    sections = []

    if resume["summary"]:
        sections.extend(
            [
                "PROFESSIONAL SUMMARY",
                resume["summary"],
            ]
        )

    if resume["skills"]:
        sections.extend(
            [
                "SKILLS",
                ", ".join(
                    resume["skills"]
                ),
            ]
        )

    section_map = [
        (
            "experience",
            "PROFESSIONAL EXPERIENCE",
        ),
        (
            "projects",
            "PROJECTS",
        ),
        (
            "education",
            "EDUCATION",
        ),
        (
            "certifications",
            "CERTIFICATIONS",
        ),
        (
            "achievements",
            "ACHIEVEMENTS",
        ),
    ]

    for key, title in section_map:

        items = resume.get(
            key,
            [],
        )

        if not items:
            continue

        sections.append(
            title
        )

        for item in items:
            sections.append(
                str(item).strip()
            )

    return "\n".join(
        item
        for item in sections
        if item
    )


# ============================================================
# DOCX GENERATION
# ============================================================

def _add_heading(
    document: Document,
    text: str,
) -> None:
    """
    Add a resume section heading.
    """

    paragraph = document.add_paragraph()

    run = paragraph.add_run(
        text
    )

    run.bold = True
    run.font.size = Pt(12)


def _add_body(
    document: Document,
    text: str,
) -> None:
    """
    Add a normal body paragraph.
    """

    paragraph = document.add_paragraph()

    run = paragraph.add_run(
        text
    )

    run.font.size = Pt(10)


def _add_bullet(
    document: Document,
    text: str,
) -> None:
    """
    Add a bullet item.
    """

    paragraph = document.add_paragraph(
        style="List Bullet"
    )

    run = paragraph.add_run(
        text
    )

    run.font.size = Pt(10)


def generate_docx_resume(
    resume: Dict[str, Any],
    output_path: str | Path,
) -> str:
    """
    Generate a standalone DOCX from a structured resume.

    This is a fallback generator. The primary application workflow
    should use update_existing_resume() to preserve the original
    document structure.
    """

    resume = normalize_resume(
        resume
    )

    output_path = Path(
        output_path
    )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    # --------------------------------------------------------
    # Summary
    # --------------------------------------------------------

    if resume["summary"]:
        _add_heading(
            document,
            "PROFESSIONAL SUMMARY",
        )

        _add_body(
            document,
            resume["summary"],
        )

    # --------------------------------------------------------
    # Skills
    # --------------------------------------------------------

    if resume["skills"]:
        _add_heading(
            document,
            "SKILLS",
        )

        _add_body(
            document,
            ", ".join(
                resume["skills"]
            ),
        )

    # --------------------------------------------------------
    # Experience
    # --------------------------------------------------------

    if resume["experience"]:
        _add_heading(
            document,
            "PROFESSIONAL EXPERIENCE",
        )

        for item in resume["experience"]:
            _add_bullet(
                document,
                item,
            )

    # --------------------------------------------------------
    # Projects
    # --------------------------------------------------------

    if resume["projects"]:
        _add_heading(
            document,
            "PROJECTS",
        )

        for item in resume["projects"]:
            _add_bullet(
                document,
                item,
            )

    # --------------------------------------------------------
    # Education
    # --------------------------------------------------------

    if resume["education"]:
        _add_heading(
            document,
            "EDUCATION",
        )

        for item in resume["education"]:
            _add_body(
                document,
                item,
            )

    # --------------------------------------------------------
    # Certifications
    # --------------------------------------------------------

    if resume["certifications"]:
        _add_heading(
            document,
            "CERTIFICATIONS",
        )

        for item in resume["certifications"]:
            _add_body(
                document,
                item,
            )

    # --------------------------------------------------------
    # Achievements
    # --------------------------------------------------------

    if resume["achievements"]:
        _add_heading(
            document,
            "ACHIEVEMENTS",
        )

        for item in resume["achievements"]:
            _add_bullet(
                document,
                item,
            )

    # --------------------------------------------------------
    # Save
    # --------------------------------------------------------

    document.save(
        str(output_path)
    )

    if not output_path.exists():
        raise RuntimeError(
            "Resume DOCX could not be created."
        )

    return str(
        output_path
    )


# ============================================================
# PLAIN TEXT FILE
# ============================================================

def generate_text_resume(
    resume: Dict[str, Any],
    output_path: str | Path,
) -> str:
    """
    Generate a plain-text resume file.
    """

    output_path = Path(
        output_path
    )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    text = resume_to_text(
        resume
    )

    output_path.write_text(
        text,
        encoding="utf-8",
    )

    return str(
        output_path
    )


# ============================================================
# RESUME PREVIEW
# ============================================================

def preview_resume(
    resume: Dict[str, Any],
) -> str:
    """
    Return a text preview of the structured resume.
    """

    return resume_to_text(
        resume
    )


# ============================================================
# BACKWARD-COMPATIBLE ALIASES
# ============================================================

def create_resume(
    resume: Dict[str, Any],
    output_path: str | Path,
) -> str:
    """
    Alias for generate_docx_resume().
    """

    return generate_docx_resume(
        resume,
        output_path,
    )


def generate_resume(
    resume: Dict[str, Any],
    output_path: str | Path,
) -> str:
    """
    Alias for generate_docx_resume().
    """

    return generate_docx_resume(
        resume,
        output_path,
    )


# ============================================================
# SCRIPT TEST
# ============================================================

if __name__ == "__main__":

    sample_resume = {
        "summary": (
            "Python developer with experience in CRM implementation "
            "and software development."
        ),
        "skills": [
            "Python",
            "SQL",
            "Git",
        ],
        "experience": [
            "Supported CRM implementation and configuration.",
            "Debugged and validated client-facing documents.",
        ],
        "projects": [
            "AI Resume Analyzer - Resume analysis and ATS matching.",
        ],
        "education": [
            "B.Sc. Statistics",
        ],
        "certifications": [],
        "achievements": [],
    }

    print(
        preview_resume(
            sample_resume
        )
    )

    print(
        "\nresume_generator.py loaded successfully."
    )