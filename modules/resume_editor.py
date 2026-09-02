"""
Structure-Preserving DOCX Resume Editor
--------------------------------------

IMPORTANT:
The user's uploaded resume is always the source document.

This editor deliberately locks structural/factual content and only
rewrites content that can be safely edited without changing the
uploaded resume's structure.

Editable:
- Professional Summary text
- Existing experience bullet text

Locked:
- Name and contact header
- Section headings
- Skills structure and category labels
- Project names and project structure
- Job titles
- Company names
- Employment dates
- Education
- Certifications
- Achievements
- Languages
- Hobbies / Additional Information
- Existing bullet count and paragraph structure

No fixed resume template is used.
"""

from pathlib import Path
import re

from docx import Document


# ============================================================
# SECTION HEADING DETECTION
# ============================================================

SECTION_HEADINGS = {
    "summary": {
        "summary",
        "professional summary",
        "profile",
        "career objective",
        "objective",
    },
    "skills": {
        "skills",
        "technical skills",
        "technical & digital skills",
        "technical and digital skills",
        "core skills",
        "core competencies",
        "key skills",
    },
    "experience": {
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "marketing experience",
        "work history",
    },
    "projects": {
        "projects",
        "project experience",
        "project experiences",
        "academic projects",
        "academic project",
        "personal projects",
        "key projects",
    },
    "education": {
        "education",
        "academic background",
        "qualifications",
    },
    "certifications": {
        "certifications",
        "certification",
        "certificates",
    },
    "achievements": {
        "achievements",
        "achievement",
        "accomplishments",
        "academic honors & extracurricular",
        "academic honors and extracurricular",
    },
    "additional": {
        "additional information",
        "additional details",
        "other information",
    },
    "languages": {
        "languages",
        "language",
    },
    "hobbies": {
        "hobbies",
        "hobbies & interests",
        "hobbies and interests",
        "interests",
    },
}


# ============================================================
# BASIC HELPERS
# ============================================================

def normalize_heading(text):
    if text is None:
        return ""

    return (
        str(text)
        .strip()
        .lower()
        .rstrip(":")
        .strip()
    )


def detect_section_heading(text):
    normalized = normalize_heading(text)

    for section_name, headings in SECTION_HEADINGS.items():
        for heading in headings:
            if normalized == normalize_heading(heading):
                return section_name

    return None


def load_docx_resume(file_path):
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"Resume file not found: {path}"
        )

    if path.suffix.lower() != ".docx":
        raise ValueError(
            "Structure-preserving editing requires a DOCX resume."
        )

    return Document(str(path))


# ============================================================
# SECTION DISCOVERY
# ============================================================

def get_section_ranges(document):
    """
    Locate existing section boundaries.

    The actual headings in the uploaded document are never changed.
    """
    sections = []

    current = None

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()

        if not text:
            continue

        detected = detect_section_heading(text)

        if detected:
            current = {
                "name": detected,
                "heading_index": index,
                "content_indexes": [],
            }
            sections.append(current)

        elif current is not None:
            current["content_indexes"].append(index)

    return sections


def get_section(document, section_name):
    for section in get_section_ranges(document):
        if section["name"] == section_name:
            return section

    return None


# ============================================================
# BULLET DETECTION
# ============================================================

def is_bullet_paragraph(paragraph):
    """
    Detect Word bullet/list paragraphs.
    """
    try:
        style_name = (
            paragraph.style.name or ""
        ).lower()

        if "bullet" in style_name:
            return True

        if "list" in style_name:
            return True
    except Exception:
        pass

    p_pr = paragraph._p.pPr

    if p_pr is not None:
        num_pr = p_pr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
        )

        if num_pr is not None:
            return True

    return False


# ============================================================
# TEXT HELPERS
# ============================================================

def _clean_generated_line(text):
    if text is None:
        return ""

    text = str(text).strip()

    for prefix in (
        "• ",
        "- ",
        "* ",
        "– ",
        "— ",
    ):
        if text.startswith(prefix):
            text = text[len(prefix):].strip()
            break

    # Remove Markdown formatting that Gemini may place inside
    # otherwise valid JSON string values.
    text = text.replace("```", "")
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)

    # Keep normal resume spacing.
    text = re.sub(r"[ \t]+", " ", text).strip()

    return text


def _flatten_generated_lines(content):
    """
    Convert AI output into plain content lines.

    Dicts are accepted only for extracting a description field.
    Structural fields such as title/company/date are intentionally
    ignored by the editor.
    """
    if content is None:
        return []

    values = []

    if isinstance(content, str):
        values = content.splitlines()

    elif isinstance(content, list):
        for item in content:
            if item is None:
                continue

            if isinstance(item, dict):
                description = ""

                for key in (
                    "description",
                    "text",
                    "content",
                    "bullet",
                ):
                    if item.get(key):
                        description = str(item[key])
                        break

                if description:
                    values.extend(
                        description.splitlines()
                    )

            else:
                values.extend(
                    str(item).splitlines()
                )

    else:
        values = [str(content)]

    result = []

    for value in values:
        cleaned = _clean_generated_line(value)

        if cleaned:
            result.append(cleaned)

    return result


def replace_paragraph_text(paragraph, new_text):
    """
    Replace text while retaining the paragraph's existing style,
    numbering, indentation and paragraph formatting.

    The first existing run is reused so its character formatting
    remains intact.
    """
    new_text = _clean_generated_line(new_text)

    if paragraph.runs:
        first_run = paragraph.runs[0]

        for run in paragraph.runs:
            run.text = ""

        first_run.text = new_text
    else:
        paragraph.add_run(new_text)


# ============================================================
# SAFE SUMMARY UPDATE
# ============================================================

def update_summary(document, summary):
    """
    Update only the first content paragraph of the existing summary.
    """
    if not summary:
        return False

    section = get_section(
        document,
        "summary",
    )

    if not section:
        return False

    paragraphs = document.paragraphs

    for index in section["content_indexes"]:
        if index >= len(paragraphs):
            continue

        paragraph = paragraphs[index]

        if paragraph.text.strip():
            replace_paragraph_text(
                paragraph,
                str(summary).strip(),
            )
            return True

    return False


# ============================================================
# LOCKED SKILLS SECTION
# ============================================================

def update_skills(document, skills):
    """
    Intentionally does nothing.

    The user's original skills structure must remain exactly as
    uploaded, including category labels and line layout.

    This prevents transformations such as:

        Programming: Java, SQL...
        Web Technologies: HTML...
        Databases: MySQL...

    becoming:

        Java
        SQL
        HTML
        MySQL

    The AI can still use the original skills for ATS analysis.
    """
    return False


# ============================================================
# EXPERIENCE
# ============================================================

def update_experience(document, experience):
    """
    Rewrite ONLY existing bullet paragraphs.

    Job title / company / dates and other non-bullet structural
    paragraphs remain untouched.

    Existing bullet count is also preserved.
    """
    if not experience:
        return False

    section = get_section(
        document,
        "experience",
    )

    if not section:
        return False

    paragraphs = document.paragraphs

    existing_bullets = []

    for index in section["content_indexes"]:
        if index >= len(paragraphs):
            continue

        paragraph = paragraphs[index]

        if is_bullet_paragraph(paragraph):
            existing_bullets.append(paragraph)

    generated_lines = _flatten_generated_lines(
        experience
    )

    if not generated_lines:
        return False

    # IMPORTANT:
    # Never add/remove bullets. Only replace the existing bullets
    # that already exist in the user's uploaded document.
    count = min(
        len(existing_bullets),
        len(generated_lines),
    )

    for index in range(count):
        replace_paragraph_text(
            existing_bullets[index],
            generated_lines[index],
        )

    return count > 0


# ============================================================
# PROJECTS
# ============================================================

def update_projects(document, projects):
    """
    Intentionally does nothing for now.

    Project paragraphs frequently contain both the project name
    and its description. Rewriting the whole paragraph can alter
    the uploaded resume's project structure.

    Therefore project content remains completely locked until the
    AI output is explicitly converted into a title/description
    structure that can be edited safely.
    """
    return False


# ============================================================
# FACTUAL SECTIONS - LOCKED
# ============================================================

def update_education(document, education):
    return False


def update_certifications(document, certifications):
    return False


def update_achievements(document, achievements):
    return False


# ============================================================
# STRUCTURE SNAPSHOT
# ============================================================

def get_structure_snapshot(document):
    """
    Capture exact existing section headings and order.
    """
    snapshot = []

    for index, paragraph in enumerate(
        document.paragraphs
    ):
        text = paragraph.text.strip()

        if not text:
            continue

        detected = detect_section_heading(text)

        if detected:
            snapshot.append(
                {
                    "index": index,
                    "internal_name": detected,
                    "heading_text": text,
                }
            )

    return snapshot


def validate_structure_preserved(
    original_document,
    updated_document,
):
    """
    Confirm that section headings and their order were preserved.
    """
    original = get_structure_snapshot(
        original_document
    )

    updated = get_structure_snapshot(
        updated_document
    )

    original_names = [
        item["internal_name"]
        for item in original
    ]

    updated_names = [
        item["internal_name"]
        for item in updated
    ]

    original_headings = [
        item["heading_text"]
        for item in original
    ]

    updated_headings = [
        item["heading_text"]
        for item in updated
    ]

    issues = []

    if original_names != updated_names:
        issues.append(
            "Section order or section set changed."
        )

    if original_headings != updated_headings:
        issues.append(
            "Existing section heading text changed."
        )

    return {
        "valid": not issues,
        "original": original,
        "updated": updated,
        "issues": issues,
    }


# ============================================================
# PROTECTED SECTION TEXT
# ============================================================

def _section_texts(document, section_name):
    """
    Return the exact paragraph texts belonging to a section.
    """
    section = get_section(
        document,
        section_name,
    )

    if not section:
        return []

    paragraphs = document.paragraphs
    values = []

    for index in section["content_indexes"]:
        if index >= len(paragraphs):
            continue

        text = paragraphs[index].text

        if text.strip():
            values.append(text)

    return values


def _snapshot_locked_content(document):
    """
    Snapshot sections that must never change.
    """
    locked_sections = (
        "skills",
        "projects",
        "education",
        "certifications",
        "achievements",
        "additional",
        "languages",
        "hobbies",
    )

    return {
        section: _section_texts(
            document,
            section,
        )
        for section in locked_sections
    }


def _validate_locked_content(
    original_document,
    updated_document,
):
    """
    Ensure locked sections remained byte-for-byte equivalent
    at paragraph-text level.
    """
    original = _snapshot_locked_content(
        original_document
    )

    updated = _snapshot_locked_content(
        updated_document
    )

    issues = []

    for section_name in original:
        if original[section_name] != updated[section_name]:
            issues.append(
                f"Locked section changed: {section_name}"
            )

    return {
        "valid": not issues,
        "issues": issues,
    }


# ============================================================
# MAIN EDITOR
# ============================================================

def update_existing_resume(
    input_path,
    output_path,
    optimized_resume,
):
    """
    Update the user's uploaded DOCX.

    This function NEVER loads a template.

    Only:
        - existing summary paragraph
        - existing experience bullet paragraphs

    may be changed.

    Everything structural/factual is locked.
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    if not isinstance(
        optimized_resume,
        dict,
    ):
        raise TypeError(
            "optimized_resume must be a dictionary."
        )

    if not input_path.exists():
        raise FileNotFoundError(
            f"Original resume not found: {input_path}"
        )

    if input_path.suffix.lower() != ".docx":
        raise ValueError(
            "Same-structure editing requires the user's original DOCX resume."
        )

    if (
        input_path.resolve()
        == output_path.resolve()
    ):
        raise ValueError(
            "The original uploaded resume cannot be overwritten."
        )

    # --------------------------------------------------------
    # Load original and create independent working document.
    # --------------------------------------------------------

    original_document = load_docx_resume(
        input_path
    )

    document = load_docx_resume(
        input_path
    )

    # --------------------------------------------------------
    # Snapshot structure and locked content.
    # --------------------------------------------------------

    original_structure = get_structure_snapshot(
        original_document
    )

    locked_before = _snapshot_locked_content(
        original_document
    )

    # --------------------------------------------------------
    # SUMMARY ONLY
    # --------------------------------------------------------

    summary = optimized_resume.get(
        "summary",
        "",
    )

    if summary:
        update_summary(
            document,
            summary,
        )

    # --------------------------------------------------------
    # SKILLS: LOCKED
    # --------------------------------------------------------

    # Do NOT call update_skills().
    # The original skills section stays exactly as uploaded.

    # --------------------------------------------------------
    # EXPERIENCE: ONLY EXISTING BULLETS
    # --------------------------------------------------------

    experience = optimized_resume.get(
        "experience",
        [],
    )

    if experience:
        update_experience(
            document,
            experience,
        )

    # --------------------------------------------------------
    # PROJECTS: LOCKED
    # --------------------------------------------------------

    # Do NOT modify projects.

    # --------------------------------------------------------
    # EDUCATION / CERTIFICATIONS / ACHIEVEMENTS:
    # LOCKED
    # --------------------------------------------------------

    # No modification.

    # --------------------------------------------------------
    # Validate section structure.
    # --------------------------------------------------------

    structure_result = validate_structure_preserved(
        original_document,
        document,
    )

    if not structure_result["valid"]:
        raise RuntimeError(
            "Structure protection failed: "
            + "; ".join(
                structure_result["issues"]
            )
        )

    # --------------------------------------------------------
    # Validate locked sections.
    # --------------------------------------------------------

    locked_after = _snapshot_locked_content(
        document
    )

    if locked_before != locked_after:
        changed = [
            section
            for section in locked_before
            if locked_before[section]
            != locked_after[section]
        ]

        raise RuntimeError(
            "Locked resume content changed: "
            + ", ".join(changed)
        )

    # --------------------------------------------------------
    # Validate no structural section was introduced/removed.
    # --------------------------------------------------------

    updated_structure = get_structure_snapshot(
        document
    )

    if original_structure != updated_structure:
        raise RuntimeError(
            "The uploaded resume's section structure was changed."
        )

    # --------------------------------------------------------
    # Save to a new file.
    # --------------------------------------------------------

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        str(output_path)
    )

    return str(output_path)


# ============================================================
# BACKWARD-COMPATIBLE ALIASES
# ============================================================

load_resume = load_docx_resume
find_section = get_section
preserve_structure = validate_structure_preserved
